from django.shortcuts import render, redirect
from .models import Store, Category, Product,Testimonial,CommandeLivraison,CategoryStore,UserPoints
from .forms import StoreForm, CategoryForm, ProductForm,RegisterForm, UpdateProfileForm,TestimonialForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.conf import settings
from django.contrib.auth import authenticate, login,logout
from django.contrib.auth.decorators import login_required
from django.views.generic import CreateView, DetailView  # DetailView pour afficher un objet, CreateView pour créer un nouvel objet
from django.urls import reverse
from .utils import get_client_ip, get_device_fingerprint


from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import RegisterForm
# def get_client_ip(request):
#     x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
#     if x_forwarded_for:
#         ip = x_forwarded_for.split(',')[0]
#     else:
#         ip = request.META.get('REMOTE_ADDR')
#     return ip

# import hashlib

# def get_device_fingerprint(request):
#     # Utiliser l'agent utilisateur pour générer un hash
#     user_agent = request.META.get('HTTP_USER_AGENT', '')
#     return hashlib.sha256(user_agent.encode()).hexdigest()

# def signup(request):
#     form = RegisterForm()

#     if request.method == 'POST':
#         form = RegisterForm(request.POST)
        
#         # Récupérer l'IP réelle de l'utilisateur
#         client_ip = get_client_ip(request)  # Obtenir l'IP réelle
#         # Récupérer l'empreinte de l'appareil
#         device_fingerprint = get_device_fingerprint(request)  # Obtenir l'empreinte de l'appareil

#         # Vérifier si un utilisateur avec la même empreinte de l'appareil ou la même IP existe
#         if CustomUser.objects.filter(device_fingerprint=device_fingerprint).exists():
#             messages.error(request, "Vous ne pouvez pas créer plusieurs comptes depuis cet appareil.")
#             return render(request, "core/signup.html", {"form": form})

#         # Vérifier si un utilisateur avec la même IP existe
#         if CustomUser.objects.filter(last_ip=client_ip).exists():
#             messages.error(request, "Votre IP réelle indique que vous avez déjà un compte.")
#             return render(request, "core/signup.html", {"form": form})

#         # Valider le formulaire, y compris le captcha
#         if form.is_valid():
#             user = form.save(commit=False)
#             user.last_ip = client_ip  # Enregistrer l'IP de l'utilisateur
#             user.device_fingerprint = device_fingerprint  # Enregistrer l'empreinte de l'appareil
#             user.save()
#             messages.success(request, "Compte créé avec succès !")
#             return redirect("signin")
#         else:
#             # Vérifier les erreurs du formulaire
#             for field, errors in form.errors.items():
#                 for error in errors:
#                     messages.error(request, f"{field}: {error}")

#     context = {"form": form}
#     return render(request, "core/signup.html", context)


# def signup(request):
#     form = RegisterForm()

#     if request.method == 'POST':
#         form = RegisterForm(request.POST)
#         if form.is_valid():
#             form.save()
#             messages.success(request, "Compte créé avec succès !")
#             return redirect("signin")
#         else:
#             # Vérifier les erreurs du formulaire
#             for field, errors in form.errors.items():
#                 for error in errors:
#                     messages.error(request, f"{field}: {error}")

#     context = {"form": form}
#     return render(request, "core/signup.html", context)
from .utils import get_client_ip, get_device_fingerprint
from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import RegisterForm
from .models import CustomUser

def signup(request):
    form = RegisterForm()

    if request.method == 'POST':
        form = RegisterForm(request.POST)
        device_fingerprint = get_device_fingerprint(request)

        # Vérifier si un utilisateur avec ce fingerprint existe déjà
        existing_device = CustomUser.objects.filter(device_fingerprint=device_fingerprint, is_active=True).exists()

        if existing_device:
            messages.error(request, "Cet appareil a déjà été utilisé pour créer un compte.")
            return render(request, "core/signup.html", {"form": form})

        if form.is_valid():
            user = form.save(commit=False)
            user.device_fingerprint = device_fingerprint
            user.last_ip = get_client_ip(request)  # Toujours garder l'IP si tu veux pour info
            user.save()
            messages.success(request, "Compte créé avec succès !")
            return redirect("signin")
        else:
            # Afficher les erreurs du formulaire
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")

    context = {"form": form}
    return render(request, "core/signup.html", context)



def signin (request):
    if request.method == 'POST':
        email = request.POST["email"]
        password= request.POST["password"]

        user = authenticate(request, email=email, password=password)
        if user is not None:
            login(request, user)
            return redirect('index')
    context= {}
    return render(request, "core/login.html", context)

def signout(request):
    logout(request)
    return redirect("index")

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

@login_required(login_url="signin")
def profile(request):
    user = request.user
    stores = Store.objects.filter(owner=user, is_active=True)
    
    # Récupérer toutes les commandes de l'utilisateur
    orders_list = Order.objects.filter(user=request.user, activated=True).order_by('-created_at')
    
    # Récupérer toutes les commandes de livraison de l'utilisateur
    commandes_list = CommandeLivraison.objects.filter(user=user).order_by('-date_commande')
    paginator = Paginator(commandes_list, 4)  # 4 commandes par page
    page_number = request.GET.get('page')
    
    try:
        commandes = paginator.page(page_number)
    except PageNotAnInteger:
        # Si la page n'est pas un entier, afficher la première page
        commandes = paginator.page(1)
    except EmptyPage:
        # Si la page est vide, afficher la dernière page
        commandes = paginator.page(paginator.num_pages)
    # Pagination - Afficher 4 commandes par page
    paginator = Paginator(orders_list, 4)  # 4 commandes par page
    page_number = request.GET.get('page')
    
    try:
        orders = paginator.page(page_number)
    except PageNotAnInteger:
        # Si la page n'est pas un entier, afficher la première page
        orders = paginator.page(1)
    except EmptyPage:
        # Si la page est vide, afficher la dernière page
        orders = paginator.page(paginator.num_pages)
     # Récupérer le nombre de points de fidélité de l'utilisateur
    try:
        user_points = UserPoints.objects.get(user=request.user)
        total_points = user_points.points  # Total des points
    except UserPoints.DoesNotExist:
        total_points = 0 

    context = {
        "user": user,
        "stores": stores,
        "orders": orders,  # Orders paginées
        "commandes": commandes,
        'total_points': total_points  # Passer les points de fidélité à la vue
    }
    
    return render(request, "core/profile.html", context)




@login_required(login_url="signin")
def update_profile(request):
    if request.user.is_authenticated:
        user = request.user
        form = UpdateProfileForm(instance=user)
        if request.method == 'POST':
            form = UpdateProfileForm(request.POST, request.FILES, instance=user)
            if form.is_valid():
                form.save()
                messages.success(request, "Profile updated successfully")
                return redirect("profile")
                
    context = {"form": form}
    return render(request, "core/update_profile.html", context)
# # Créer une boutique

# def create_store(request):
#     if request.method == 'POST':
#         form = StoreForm(request.POST)
#         if form.is_valid():
#             store = form.save(commit=False)
#             store.owner = request.user
#             store.save()
#             return redirect('shop:store_detail', store_id=store.id)
#     else:
#         form = StoreForm()
#     return render(request, 'core/create_store.html', {'form': form})
@login_required
def create_store(request):
    if request.method == 'POST':
        form = StoreForm(request.POST, request.FILES, initial={'user': request.user})
        if form.is_valid():
            store = form.save(commit=False)
            store.owner = request.user  # Associer l'utilisateur à ce store
            store.save()

            messages.success(request, "Votre store a été créé avec succès, mais il n'est pas encore activé pour être visible. Nous vous prions de contacter notre administration pour l'activer.")
            return redirect('profile')  # Redirection après succès
    else:
        form = StoreForm(initial={'user': request.user})

    return render(request, 'core/create_store.html', {'form': form})

# AJAX
def load_cellules(request):
    typestore_id = request.GET.get('typestore_id')
    categorystores = CategoryStore.objects.filter(typestore_id=typestore_id).all()
    return render(request, 'core/cellule_dropdown_list_options.html', {'categorystores': categorystores})



def get_cellules(request):
    typestore_id = request.GET.get('typestore')
    categorystores = CategoryStore.objects.filter(typestore_id=typestore_id).values('id', 'nom')
    return JsonResponse(list(categorystores), safe=False)




@login_required
def edit_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)  # Lookup by slug instead of ID
    if request.method == 'POST':
        form = StoreForm(request.POST, request.FILES, instance=store)
        if form.is_valid():
            form.save()
            messages.success(request, "Le store a été mis à jour avec succès !")
            return redirect('profile')  # Or wherever you want to redirect after editing
    else:
        form = StoreForm(instance=store)

    return render(request, 'core/edit_store.html', {'form': form, 'store': store})

# delete store
@login_required
def delete_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)  # Lookup by slug instead of ID
    if request.method == 'POST':
        store.delete()
        messages.success(request, "Le store a été supprimé avec succès !")
        return redirect('profile')  # Or wherever you want to redirect after deletion
    return render(request, 'core/delete_store.html', {'store': store})


from django.shortcuts import render
from django.http import JsonResponse
from .models import Store, CategoryStore, Typestore
from geopy.distance import geodesic
from django.core.paginator import Paginator
from geopy.distance import geodesic
from django.core.paginator import Paginator
from django.shortcuts import render
from .models import Store, CategoryStore, Typestore
from django.shortcuts import render
from geopy.distance import geodesic
from .models import Store, CategoryStore, Typestore
from django.core.paginator import Paginator
from django.shortcuts import render, redirect
from django.http import JsonResponse
from .models import Store, UserLocation
from geopy.distance import geodesic

def save_user_location(request):
    # Récupérer les coordonnées de l'utilisateur depuis la requête
    user_lat = request.GET.get('lat')
    user_lon = request.GET.get('lon')

    if user_lat and user_lon:
        # Créer ou mettre à jour la localisation de l'utilisateur
        user_location, created = UserLocation.objects.update_or_create(
            user=request.user,
            defaults={'latitude': user_lat, 'longitude': user_lon}
        )

        return JsonResponse({'message': 'Location saved successfully'})

    return JsonResponse({'message': 'Failed to save location'}, status=400)


from django.shortcuts import render, redirect
from geopy.distance import geodesic
from .models import Store, UserLocation
from django.core.paginator import Paginator

from django.shortcuts import render, redirect
from geopy.distance import geodesic
from .models import Store, UserLocation
from django.core.paginator import Paginator
from django.contrib import messages

from django.shortcuts import render
from django.core.paginator import Paginator
from geopy.distance import geodesic
from .models import Store, UserLocation
from django.contrib import messages
from geopy.distance import geodesic
from django.core.paginator import Paginator
from .models import Store, UserLocation, CategoryStore, Typestore

from geopy.distance import geodesic
from django.core.paginator import Paginator
from django.contrib import messages
@login_required(login_url="signin")
def search_stores(request):
    # On vérifie si le formulaire a été envoyé via POST
    if request.method == "POST":
        user_lat = request.POST.get('lat')
        user_lon = request.POST.get('lon')

        if user_lat and user_lon:
            try:
                # Convertir les coordonnées en flottant
                user_lat = float(user_lat)
                user_lon = float(user_lon)

                # Enregistrer ou mettre à jour la localisation de l'utilisateur
                user_location, created = UserLocation.objects.get_or_create(user=request.user)
                user_location.latitude = user_lat
                user_location.longitude = user_lon
                user_location.save()

                # Afficher un message de succès
                messages.success(request, "Votre position a été enregistrée avec succès. Recherche en cours...")

            except ValueError:
                messages.error(request, "Erreur lors de l'enregistrement de la position.")
                return redirect('search_stores')

    # Récupérer les données de localisation et les magasins
    user_location = UserLocation.objects.filter(user=request.user).first()
    stores = Store.objects.all()

    # Appliquer les filtres si présents
    category_id = request.POST.get('category')
    typestore_id = request.POST.get('typestore')
    store_name = request.POST.get('store_name', '').strip()
    store_address = request.POST.get('store_address', '').strip()

    if category_id:
        stores = stores.filter(categorystore_id=category_id)

    if typestore_id:
        stores = stores.filter(typestore_id=typestore_id)

    if store_name:
       store_name = ' '.join(store_name.split())  # Normaliser l'espace
    stores = stores.filter(name__icontains=store_name)  # Recherche par nom

    
    if store_address:
       store_address = ' '.join(store_address.split())  # Normaliser l'espace
    stores = stores.filter(adresse__icontains=store_address)  # Utiliser "adresse" et non "address"


    # Initialiser la liste des magasins proches
    nearby_stores = []

    # Filtrer par proximité si la position de l'utilisateur est disponible
    if user_location:
        user_coords = (float(user_location.latitude), float(user_location.longitude))

        for store in stores:
            if store.latitude and store.longitude:
                store_coords = (store.latitude, store.longitude)
                distance = geodesic(user_coords, store_coords).km
                if distance <= 5:  # Rayon de 5 km pour les magasins proches
                    store.distance = distance
                    nearby_stores.append(store)
        
        # Trier les magasins par distance (les plus proches en premier)
        nearby_stores.sort(key=lambda x: x.distance)

        # Pagination des résultats
        paginator = Paginator(nearby_stores, 6)  # Affichage de 1 magasin par page
        page_number = request.GET.get('page')
        stores = paginator.get_page(page_number)
    else:
        stores = []  # Si pas de localisation, afficher aucun magasin

    # Récupérer les catégories et types pour les afficher dans le formulaire
    categories = CategoryStore.objects.all()
    types = Typestore.objects.all()

    # Calculer le nombre de magasins proches
    num_stores = len(nearby_stores)  # Le nombre de magasins proches

    return render(request, 'core/search_stores.html', {
        'stores': stores,
        'categories': categories,
        'types': types,
        'num_stores': num_stores,  # Passer le nombre de magasins au template
        'store_address': store_address,
    })



from geopy.distance import geodesic
from django.core.paginator import Paginator
from django.shortcuts import render
from .models import Product, UserLocation

from geopy.distance import geodesic
from geopy.distance import geodesic
from django.contrib import messages
from django.core.paginator import Paginator
@login_required(login_url="signin")
def search_products(request):
    # Vérification si le formulaire est envoyé en POST
    if request.method == "POST":
        user_lat = request.POST.get('lat')
        user_lon = request.POST.get('lon')

        if user_lat and user_lon:
            try:
                # Convertir la latitude et longitude en float
                user_lat = float(user_lat)
                user_lon = float(user_lon)

                # Enregistrer ou mettre à jour la localisation de l'utilisateur
                user_location, created = UserLocation.objects.get_or_create(user=request.user)
                user_location.latitude = user_lat
                user_location.longitude = user_lon
                user_location.save()

                # Message de succès
                messages.success(request, "Votre position a été enregistrée avec succès. Recherche en cours...")

            except ValueError:
                messages.error(request, "Erreur lors de l'enregistrement de la position.")
                return redirect('search_products')

    # Récupérer la localisation de l'utilisateur et les produits
    user_location = UserLocation.objects.filter(user=request.user).first()
    products = Product.objects.all()

    # Appliquer les filtres pour le nom et la catégorie
    product_name = request.POST.get('product_name', '').strip()
    category_id = request.POST.get('category')

    if product_name:
        product_name = ' '.join(product_name.split())  # Normaliser l'espace
        products = products.filter(name__icontains=product_name)

    if category_id:
        products = products.filter(category_id=category_id)

    # Vérifie que l'utilisateur a une localisation
    if user_location:
        user_coords = (float(user_location.latitude), float(user_location.longitude))
        nearby_products = []  # Liste pour stocker les produits proches

        for product in products:
            store = product.store  # Récupérer le store auquel appartient le produit
            if store.latitude and store.longitude:
                store_coords = (store.latitude, store.longitude)
                distance = geodesic(user_coords, store_coords).km  # Calculer la distance entre l'utilisateur et le store
                if distance <= 5:  # Produits dont le store est dans un rayon de 5 km
                    product.store.distance = distance  # Ajouter la distance du store à l'objet store
                    nearby_products.append(product)

        # Trier les produits par distance (les plus proches en premier)
        nearby_products.sort(key=lambda x: x.store.distance)

        # Pagination des produits
        paginator = Paginator(nearby_products, 6)  # 6 produits par page
        page_number = request.GET.get('page')
        products = paginator.get_page(page_number)

        # Compter le nombre de produits proches
        num_products = len(nearby_products)
    else:
        products = []
        num_products = 0  # Aucun produit trouvé si la localisation n'est pas disponible

    # Récupérer les catégories de produits pour le filtrage
    categories = Category.objects.all()

    return render(request, 'core/search_products.html', {
        'products': products,
        'categories': categories,
        'num_products': num_products,  # Passer le nombre de produits au template
    })

# def search_products(request):
#     # Vérification si le formulaire est envoyé en POST
#     if request.method == "POST":
#         user_lat = request.POST.get('lat')
#         user_lon = request.POST.get('lon')

#         if user_lat and user_lon:
#             try:
#                 # Convertir la latitude et longitude en float
#                 user_lat = float(user_lat)
#                 user_lon = float(user_lon)

#                 # Enregistrer ou mettre à jour la localisation de l'utilisateur
#                 user_location, created = UserLocation.objects.get_or_create(user=request.user)
#                 user_location.latitude = user_lat
#                 user_location.longitude = user_lon
#                 user_location.save()

#                 # Message de succès
#                 messages.success(request, "Votre position a été enregistrée avec succès. Recherche en cours...")

#             except ValueError:
#                 messages.error(request, "Erreur lors de l'enregistrement de la position.")
#                 return redirect('search_products')

#     # Récupérer la localisation de l'utilisateur et les produits
#     user_location = UserLocation.objects.filter(user=request.user).first()
#     products = Product.objects.all()

#     # Appliquer les filtres pour le nom et la catégorie
#     product_name = request.POST.get('product_name', '').strip()
#     category_id = request.POST.get('category')

#     if product_name:
#        product_name = ' '.join(product_name.split())  # Normaliser l'espace
#        products = products.filter(name__icontains=product_name)

#     if category_id:
#         products = products.filter(category_id=category_id)

#     # Vérifie que l'utilisateur a une localisation
#     if user_location:
#         user_coords = (float(user_location.latitude), float(user_location.longitude))
#         nearby_products = []

#         for product in products:
#             store = product.store  # Récupérer le store auquel appartient le produit
#             if store.latitude and store.longitude:
#                 store_coords = (store.latitude, store.longitude)
#                 distance = geodesic(user_coords, store_coords).km  # Calculer la distance entre l'utilisateur et le store
#                 if distance <= 5:  # Produits dont le store est dans un rayon de 5 km
#                     product.store.distance = distance  # Ajouter la distance du store à l'objet store
#                     nearby_products.append(product)

#         # Trier les produits par distance (les plus proches en premier)
#         nearby_products.sort(key=lambda x: x.store.distance)

#         # Pagination des produits
#         paginator = Paginator(nearby_products, 6)  # 6 produits par page
#         page_number = request.GET.get('page')
#         products = paginator.get_page(page_number)

#         # Compter le nombre de produits proches
#         num_products = len(nearby_products)
#     else:
#         products = []
#         num_products = 0  # Aucun produit trouvé si la localisation n'est pas disponible

#     # Récupérer les catégories de produits pour le filtrage
#     categories = Category.objects.all()

#     return render(request, 'core/search_products.html', {
#         'products': products,
#         'categories': categories,
#         'num_products': num_products,  # Passer le nombre de produits au template
#     })




# def search_stores(request):
#     # On vérifie si le formulaire a été envoyé via POST
#     if request.method == "POST":
#         user_lat = request.POST.get('lat')
#         user_lon = request.POST.get('lon')

#         if user_lat and user_lon:
#             try:
#                 # Convertir les coordonnées en flottant
#                 user_lat = float(user_lat)
#                 user_lon = float(user_lon)

#                 # Enregistrer ou mettre à jour la localisation de l'utilisateur
#                 user_location, created = UserLocation.objects.get_or_create(user=request.user)
#                 user_location.latitude = user_lat
#                 user_location.longitude = user_lon
#                 user_location.save()

#                 # Afficher un message de succès
#                 messages.success(request, "Votre position a été enregistrée avec succès. Recherche en cours...")

#             except ValueError:
#                 messages.error(request, "Erreur lors de l'enregistrement de la position.")
#                 return redirect('search_stores')

#     # Récupérer la localisation de l'utilisateur
#     user_location = UserLocation.objects.filter(user=request.user).first()

#     stores = Store.objects.all()

#     if user_location:
#         user_coords = (float(user_location.latitude), float(user_location.longitude))
#         nearby_stores = []

#         # Calculer la distance entre l'utilisateur et les magasins
#         for store in stores:
#             if store.latitude and store.longitude:
#                 store_coords = (store.latitude, store.longitude)
#                 distance = geodesic(user_coords, store_coords).km
#                 if distance <= 5:  # Rayon de 5 km pour les magasins proches
#                     store.distance = distance
#                     nearby_stores.append(store)

#         # Pagination des résultats
#         paginator = Paginator(nearby_stores, 9)  # Affichage de 9 magasins par page
#         page_number = request.GET.get('page')
#         stores = paginator.get_page(page_number)
#     else:
#         stores = []

#     return render(request, 'core/search_stores.html', {
#         'stores': stores,
#     })






from geopy.distance import geodesic
from django.http import JsonResponse
from .models import Store, CategoryStore, Typestore

def get_stores_nearby(request):
    # Récupérer la latitude et la longitude de l'utilisateur
    user_lat = request.GET.get('lat')
    user_lon = request.GET.get('lon')

    # Récupérer les filtres de catégorie et de type de magasin
    category_id = request.GET.get('category')
    typestore_id = request.GET.get('typestore')

    stores = Store.objects.all()

    # Appliquer les filtres de catégorie et de type de magasin
    if category_id:
        stores = stores.filter(categorystore_id=category_id)
    if typestore_id:
        stores = stores.filter(typestore_id=typestore_id)

    # Calculer la distance et trouver les magasins proches
    nearby_stores = []
    if user_lat and user_lon:
        user_coords = (float(user_lat), float(user_lon))

        # Calculer la distance pour chaque magasin
        for store in stores:
            if store.latitude and store.longitude:
                store_coords = (store.latitude, store.longitude)
                distance = geodesic(user_coords, store_coords).km
                if distance <= 5:  # Filtrer par rayon de 5 km
                    nearby_stores.append({
                        'name': store.name,
                        'category': store.categorystore.nom,
                        'address': store.adresse,
                        'distance': distance,
                        'thumbnail': store.thumbnail.url if store.thumbnail else None,
                        'slug': store.slug
                    })

    # Retourner les magasins proches sous forme de JSON
    return JsonResponse({'stores': nearby_stores})

# Créer une catégorie
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Category, Store
from .forms import CategoryForm
from django.http import HttpResponseForbidden


# @login_required
# def create_category(request, slug):
#     # Get the store using the slug and ensure the logged-in user is the owner
#     store = get_object_or_404(Store, slug=slug, owner=request.user)
    
#     # Handle category creation
#     if request.method == 'POST' and 'create_category' in request.POST:
#         form = CategoryForm(request.POST)
#         if form.is_valid():
#             category = form.save(commit=False)
#             category.store = store
#             category.save()
#             messages.success(request, "Catégorie créée avec succès!")
#             return redirect('store_detail', slug=store.slug)  # Refresh the page to show the new category
#     else:
#         form = CategoryForm()

#     # Get all categories for this store
#     categories = Category.objects.filter(store=store)

#     return render(request, 'core/create_category.html', {'form': form, 'store': store, 'categories': categories})
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from .models import Category, Store
from .forms import CategoryForm

@login_required
def create_category(request, store_id):
    # Récupérer le store via son id et s'assurer que l'utilisateur est le propriétaire
    store = get_object_or_404(Store, id=store_id, owner=request.user)

    # Créer une nouvelle catégorie
    if request.method == 'POST' and 'create_category' in request.POST:
        form = CategoryForm(request.POST)
        if form.is_valid():
            category = form.save(commit=False)
            category.store = store
            category.save()
            messages.success(request, "Catégorie créée avec succès!")
            return redirect('create_category', store_id=store.id)  # Redirection pour éviter une soumission multiple
    else:
        form = CategoryForm()

    # Gérer la suppression d'une catégorie
    if request.method == 'POST' and 'delete_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        category.delete()
        messages.success(request, "Catégorie supprimée avec succès!")
        return redirect('create_category', store_id=store.id)

    # Gérer la modification d'une catégorie
    if request.method == 'POST' and 'edit_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        form = CategoryForm(request.POST, instance=category)
        if form.is_valid():
            form.save()
            messages.success(request, "Catégorie mise à jour avec succès!")
            return redirect('create_category', store_id=store.id)

    # Récupérer toutes les catégories existantes pour le store
    categories = Category.objects.filter(store=store).order_by('-created_at')

    return render(request, 'core/create_category.html', {
        'form': form,
        'store': store,
        'categories': categories
    })

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from .models import Category, Store

@login_required
def list_categories(request, store_id):
    # Récupérer le store via son ID et s'assurer que l'utilisateur est le propriétaire
    store = get_object_or_404(Store, id=store_id, owner=request.user)

    # Supprimer une catégorie
    if request.method == 'POST' and 'delete_category' in request.POST:
        category_id = request.POST.get('category_id')
        category = get_object_or_404(Category, id=category_id, store=store)
        category.delete()
        messages.success(request, "Catégorie supprimée avec succès!")
        return redirect('list_categories', store_id=store.id)

    # Récupérer toutes les catégories pour le store
    categories = Category.objects.filter(store=store)

    return render(request, 'core/list_categories.html', {
        'store': store,
        'categories': categories,
    })


from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from .models import Category
from .forms import CategoryForm

@login_required
def edit_category(request, slug):
    # Récupérer le store auquel les catégories appartiennent
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    
    # Récupérer toutes les catégories du store
    categories = Category.objects.filter(store=store)

    if request.method == 'POST':
        # Si un formulaire de modification ou de suppression est soumis
        if 'edit_category' in request.POST:
            category_id = request.POST.get('category_id')
            category = get_object_or_404(Category, id=category_id, store=store)
            form = CategoryForm(request.POST, instance=category)
            if form.is_valid():
                form.save()
                messages.success(request, "Catégorie modifiée avec succès!")
                return redirect('edit_category', slug=store.slug)
        
        elif 'delete_category' in request.POST:
            category_id = request.POST.get('category_id')
            category = get_object_or_404(Category, id=category_id, store=store)
            category.delete()
            messages.success(request, "Catégorie supprimée avec succès!")
            return redirect('edit_category', slug=store.slug)

    return render(request, 'core/edit_category.html', {'store': store, 'categories': categories})

# Créer un produit

from decimal import Decimal

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Store
from .forms import ProductForm

# views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Store, Product, Category,AssignerCategory
from .forms import ProductForm
# views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Store, Product, Category
from .forms import ProductForm,AssignerCategoryForm
import logging
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Store, Product
from .forms import ProductForm
@login_required
def create_product(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            product = form.save(commit=False)
            product.store = store
            product.category = None  # Pas de catégorie par défaut
            product.save()

            messages.success(request, "Produit créé avec succès ! Veuillez maintenant l'assigner à une catégorie.")

            print(f"✅ Produit {product.id} créé. Redirection vers assign_category.")  # Debug

            return redirect('assign_category', product_id=product.id)  # ✅ Redirection ici

        else:
            print("❌ Formulaire invalide :", form.errors)  # Debug en cas d'erreur

    else:
        form = ProductForm()

    return render(request, 'core/create_product.html', {'form': form, 'store': store})


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Store, Category, Product, AssignerCategory

@login_required
def assign_category(request, store_id):
    store = get_object_or_404(Store, id=store_id, owner=request.user)
    
    # Récupérer catégories et produits du store (triés du plus récent au plus ancien)
    categories = store.categories.all().order_by('-id')
    products = store.products.filter(category=None).order_by('-id')  # Seulement ceux sans catégorie

    if request.method == 'POST':
        category_id = request.POST.get('category')
        product_id = request.POST.get('product')

        if category_id and product_id:
            category = get_object_or_404(Category, id=category_id, store=store)
            product = get_object_or_404(Product, id=product_id, store=store)

            # Créer l'assignation
            AssignerCategory.objects.create(product=product, category=category)
            product.category = category  # Met à jour la catégorie du produit
            product.save()

            messages.success(request, f"Catégorie '{category.name}' assignée à '{product.name}' avec succès !")
            return redirect('store_detail', slug=store.slug)
  

    return render(request, 'core/assign_category.html', {'store': store, 'categories': categories, 'products': products})





# @login_required
# def create_product(request, store_id):
#     store = Store.objects.get(id=store_id, owner=request.user)  # Récupérer le store

#     if request.method == 'POST':
#         form = ProductForm(request.POST, request.FILES)
#         if form.is_valid():
#             product = form.save(commit=False)  # Créer l'objet Product mais ne le sauvegarde pas encore
#             product.store = store  # Associer ce produit au store
#             product.save()  # Sauvegarde le produit avec le prix incluant la commission

#             messages.success(request, "Votre produit a été créé avec succès !")
#             return redirect('store_detail', store_id=store.id)  # Redirection après succès
#     else:
#         form = ProductForm()

#     return render(request, 'core/create_product.html', {'form': form, 'store': store})


# Détails d'une boutique
# def store_detail(request, store_id):
#     store = Store.objects.get(id=store_id)
#     return render(request, 'core/store_detail.html', {'store': store})

from django.shortcuts import render, redirect, get_object_or_404
from .models import Product, Cart, CartItem
from django.contrib.auth.decorators import login_required
from django.http import Http404

# Fonction utilitaire pour obtenir ou créer un panier
def get_or_create_cart(user):
    cart, created = Cart.objects.get_or_create(user=user, is_ordered=False)
    return cart

# Ajouter un produit au panier
from django.http import JsonResponse

# @login_required
# def add_to_cart_ajax(request, product_id):
#     product = get_object_or_404(Product, id=product_id)
#     cart = get_or_create_cart(request.user)

#     # Vérifier si l'article est déjà dans le panier
#     cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
#     if not created:
#         cart_item.quantity += 1
#     else:
#         cart_item.quantity = 1
#     cart_item.save()

#     # Calculer le nombre total d'articles et le total du panier
#     total_items = cart.get_item_count()
#     total_price = cart.get_total()

#     # Retourner ces informations sous forme de réponse JSON
#     return JsonResponse({'total_items': total_items, 'total_price': total_price})
from .models import Cart, CartItem

# @login_required
# def add_to_cart_ajax(request, product_id):
#     product = get_object_or_404(Product, id=product_id)
#     cart = get_or_create_cart(request.user)

#     # Vérifier si l'article est déjà dans le panier
#     cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
#     if not created:
#         cart_item.quantity += 1
#     else:
#         cart_item.quantity = 1
#     cart_item.save()

#     # Calculer le nombre total d'articles et le total du panier
#     total_items = cart.get_item_count()
#     total_price = cart.get_total()

#     # Retourner ces informations sous forme de réponse JSON
#     return JsonResponse({'total_items': total_items, 'total_price': total_price})
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.urls import reverse

# @login_required
# def add_to_cart_ajax(request, product_id):
#     product = get_object_or_404(Product, id=product_id)
#     cart = get_or_create_cart(request.user)

#     # Vérifier si l'article est déjà dans le panier
#     cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
#     if not created:
#         cart_item.quantity += 1
#     else:
#         cart_item.quantity = 1
#     cart_item.save()

#     # Calculer le nombre total d'articles et le total du panier
#     total_items = cart.get_item_count()
#     total_price = cart.get_total()

#     # Rediriger vers la page de confirmation
#     return redirect('add_to_cart_success')  # Redirige vers la page de succès
@login_required(login_url="signin")
def add_to_cart_ajax(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    cart = get_or_create_cart(request.user)

    # Vérifier si l'article est déjà dans le panier
    cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
    if not created:
        cart_item.quantity += 1
    else:
        cart_item.quantity = 1
    cart_item.save()

    # Calculer le nombre total d'articles dans le panier
    total_items = cart.get_item_count()

    # Retourner la redirection mais aussi mettre à jour dynamiquement la page de confirmation
    return redirect('add_to_cart_success')


from django.http import JsonResponse
from .models import Cart

@login_required
def get_cart_items_count(request):
    cart = get_or_create_cart(request.user)
    total_items = cart.get_item_count()
    return JsonResponse({
        'total_items': total_items,
    })

def context_processors(request):
    if request.user.is_authenticated:
        cart = get_or_create_cart(request.user)
    else:
        cart = None
    return {'cart': cart}

def get_or_create_cart(user):
    """
    Cette fonction vérifie si l'utilisateur a déjà un panier actif.
    Si oui, il le retourne. Sinon, il crée un nouveau panier.
    """
    cart, created = Cart.objects.get_or_create(user=user, is_active=True)
    return cart


from django.shortcuts import render

@login_required(login_url="signin")
def add_to_cart_success(request):
    # Obtenez le panier de l'utilisateur
    cart = get_or_create_cart(request.user)
    total_items = cart.get_item_count()
    total_price = cart.get_total()

    # Vérifie s'il y a un produit dans le panier
    last_added_product = cart.items.last().product if cart.items.exists() else None

    if last_added_product is None:
        return redirect('index')  # Redirige vers la liste des produits si le panier est vide

    return render(request, 'core/add_to_cart_success.html', {
        'product': last_added_product,
        'total_items': total_items,
        'total_price': total_price,
    })



# Voir le contenu du panier
from .models import Cart

@login_required(login_url="signin")
def cart_detail(request):
    # Récupère ou crée le panier de l'utilisateur connecté
    cart = get_or_create_cart(request.user)
    
    context = {
        'cart': cart
    }
    
    return render(request, 'core/cart_detail.html', context)
# @login_required
# def cart_detail(request):
#     # Récupère ou crée le panier de l'utilisateur connecté
#     cart = get_or_create_cart(request.user)

#     # Crée la commande à partir du panier si elle n'existe pas
#     if not cart.is_ordered:
#         order = Order.objects.create(
#             user=request.user,
#             store=cart.items.first().product.store,  # Associe la commande au premier magasin dans le panier
#         )
#         for item in cart.items.all():
#             OrderItem.objects.create(
#                 order=order,
#                 product=item.product,
#                 quantity=item.quantity,
#                 price_at_time_of_order=item.product.price
#             )
#         order.calculate_total()
#         cart.is_ordered = True
#         cart.save()
#     else:
#         # Si le panier est déjà commandé, récupère la commande existante
#         order = Order.objects.filter(user=request.user, status='pending').first()

#     context = {
#         'cart': cart,
#         'order': order  # Assure-toi de passer `order` dans le contexte
#     }

#     return render(request, 'core/cart_detail.html', context)


# Modifier la quantité d'un produit dans le panier
@login_required
def update_cart(request, cart_item_id, quantity):
    # Récupérer l'élément du panier correspondant
    cart_item = get_object_or_404(CartItem, id=cart_item_id, cart__user=request.user)

    # Si la quantité vient du formulaire POST, l'utiliser à la place de celle dans l'URL
    if request.method == "POST":
        quantity = int(request.POST.get('quantity', cart_item.quantity))  # Si pas de quantity dans le POST, conserver l'ancienne

    # Si la quantité est inférieure ou égale à zéro, supprimer l'élément du panier
    if quantity <= 0:
        cart_item.delete()
    else:
        # Sinon, mettre à jour la quantité
        cart_item.quantity = quantity
        cart_item.save()
    
    # Rediriger vers la vue du panier
    return redirect('cart_detail')



@login_required
def remove_from_cart(request, cart_item_id):
    try:
        # Récupérer l'élément du panier correspondant à l'utilisateur
        cart_item = CartItem.objects.get(id=cart_item_id, cart__user=request.user)
    except CartItem.DoesNotExist:
        # Si l'élément n'existe pas (par exemple, si le panier a été vidé), rediriger vers le panier
        return redirect('cart_detail')

    # Supprimer l'élément du panier
    cart_item.delete()
    
    # Rediriger vers la vue du panier après la suppression
    return redirect('cart_detail')


# # Supprimer un produit du panier
# @login_required
# def remove_from_cart(request, cart_item_id):
#     # Récupérer l'élément du panier correspondant
#     cart_item = get_object_or_404(CartItem, id=cart_item_id, cart__user=request.user)
    
#     # Supprimer l'élément du panier
#     cart_item.delete()
    
#     # Rediriger vers la vue du panier
#     return redirect('cart_detail')


from django.contrib.auth.mixins import LoginRequiredMixin

class TestimonialCreateView(LoginRequiredMixin, CreateView):
    model = Testimonial
    form_class = TestimonialForm
    template_name = 'testimonial_form.html'  # Créez ce template pour le formulaire
    context_object_name = 'form'

    def form_valid(self, form):
        form.instance.user = self.request.user  # L'utilisateur actuel est assigné au témoignage
        form.instance.store = Store.objects.get(id=self.kwargs['pk'])  # Associe le témoignage au magasin
        return super().form_valid(form)

    def get_success_url(self):
        store_id = self.kwargs['pk']
        return reverse('store_detail', kwargs={'pk': store_id}) 
         
from django.conf import settings
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from core.models import Cart

# @login_required
# def checkout(request):
#     # Récupérer ou créer le panier de l'utilisateur
#     cart = get_or_create_cart(request.user)
    
#     # Calculer le total du panier et le diviser par 100 pour obtenir la valeur en dollars
#     total_price_in_cents = cart.get_total()
#     total_price_in_dollars = total_price_in_cents / 100  # Conversion en dollars

#     # Passer les valeurs au template
#     pub_key = settings.STRIPE_API_KEY_PUBLISHABLE

#     return render(request, 'core/checkout.html', {
#         'pub_key': pub_key,
#         'cart': cart,
#         'total_price_in_dollars': total_price_in_dollars
#     })


# def success(request):
#     return render(request, 'core/success.html',)

from django.shortcuts import render, redirect
from .models import Cart, CartItem, Order, OrderItem, Product
from django.contrib.auth.decorators import login_required
from django.http import Http404
from django.shortcuts import render, redirect
from .models import Order, OrderItem, Cart
from django.contrib.auth.decorators import login_required
from django.db.models import Sum

@login_required
def payment_success(request):
    # Récupérer le panier de l'utilisateur
    cart = get_or_create_cart(request.user)
   
    if cart.get_item_count() == 0:
        return redirect('cart_detail')
   
    # Créer une commande pour l'utilisateur
    order = Order.objects.create(
        user=request.user,
        status='paid',
    )
   
    # Ajouter les articles de la commande
    for cart_item in cart.items.all():
        OrderItem.objects.create(
            order=order,
            product=cart_item.product,
            quantity=cart_item.quantity,
            price_at_time_of_order=cart_item.product.price,
            store=cart_item.product.store  # Associer chaque article au magasin du produit
        )
   
    # Mettre à jour le total de la commande
    order.calculate_total()

    # Vider le panier de l'utilisateur après la commande
    cart.items.all().delete()

    # Retourner à la page de confirmation avec la commande
    return render(request, 'core/payment_success.html', {'order': order})


# @login_required
# def payment_success(request):
#     # Récupérer le panier de l'utilisateur
#     cart = get_or_create_cart(request.user)
   
#     if cart.get_item_count() == 0:
#         return redirect('cart_detail')
   
#     # Créer une commande
#     order = Order.objects.create(
#         user=request.user,
#         store=cart.items.first().product.store,  # Utilise le store du premier produit du panier
#         status='paid',
#     )
   
#     # Ajouter les articles de la commande
#     for cart_item in cart.items.all():
#         OrderItem.objects.create(
#             order=order,
#             product=cart_item.product,
#             quantity=cart_item.quantity,
#             price_at_time_of_order=cart_item.product.price
#         )

#     # Mettre à jour le total de la commande
#     order.calculate_total()

#     # Vider le panier de l'utilisateur après la commande
#     cart.items.all().delete()

#     # Retourner à la page de confirmation
#     return render(request, 'core/payment_success.html', {'order': order})

# payement mobile money

@login_required
def mobile_money_payment_success(request):
    # Récupérer le paiement Mobile Money validé
    payment = MobileMoneyPayment.objects.filter(user=request.user, activated=True).last()
    
    if not payment:
        return redirect('mobile_money_checkout')  # Si aucun paiement validé, rediriger

    # Créer une commande pour Mobile Money
    cart = get_or_create_cart(request.user)
    order = Order.objects.create(
        user=request.user,
        store=cart.items.first().product.store,
        status='paid',
    )

    # Ajouter les articles de la commande
    for cart_item in cart.items.all():
        OrderItem.objects.create(
            order=order,
            product=cart_item.product,
            quantity=cart_item.quantity,
            price_at_time_of_order=cart_item.product.price
        )

    order.calculate_total()  # Mettre à jour le total de la commande

    # Vider le panier de l'utilisateur après la commande
    cart.items.all().delete()

    return render(request, 'core/mobile_money_payment_success.html', {'order': order})
# views.py
@login_required
def mobile_money_waiting(request, payment_id):
    payment = get_object_or_404(MobileMoneyPayment, id=payment_id, user=request.user)

    if payment.status == 'validated':  # Si le paiement est validé, on redirige vers la page de succès
        return redirect('mobile_money_payment_success')

    return render(request, 'core/mobile_money_waiting.html', {'payment': payment})


from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .forms import MobileMoneyPaymentForm
from .models import MobileMoneyPayment, Order, OrderItem

@login_required
def mobile_money_checkout(request):
    # Vérifier si le panier est vide
    cart = get_or_create_cart(request.user)
    if cart.get_item_count() == 0:
        return redirect('cart_detail')  # ou 'core:cart_detail' si bien défini
  # Si le panier est vide, on redirige

    # Calculer le total des articles et le montant total
    total_items = cart.get_item_count()
    total_amount = sum(item.product.price_with_commission * item.quantity for item in cart.items.all())

    if request.method == 'POST':
        form = MobileMoneyPaymentForm(request.POST)
        if form.is_valid():
            # Sauvegarder le paiement Mobile Money
            payment = form.save(commit=False)
            payment.user = request.user
            payment.status = 'pending'  # Statut initial en attente
            payment.save()

            # Créer l'ordre mais sans passer de store à Order() pour éviter l'erreur
            order = Order.objects.create(
                user=request.user,
                status='pending',
                activated=False  # Désactivé par défaut pour Mobile Money
            )

            # Ajouter les articles de la commande
            for cart_item in cart.items.all():
                OrderItem.objects.create(
                    order=order,
                    product=cart_item.product,
                    quantity=cart_item.quantity,
                    price_at_time_of_order=cart_item.product.price_with_commission
                )

            # Calculer le total de la commande
            order.calculate_total()

            # Vider le panier de l'utilisateur après la commande
            cart.items.all().delete()  # Suppression des articles dans le panier

            # Message de succès et redirection vers la page d'attente
            messages.success(request, 'Votre paiement Mobile Money a été soumis avec succès. Veuillez attendre la validation de l\'admin.')
            return redirect('mobile_money_waiting', payment_id=payment.id)
    else:
        form = MobileMoneyPaymentForm()

    return render(request, 'core/mobile_money_checkout.html', {
        'form': form,
        'total_items': total_items,
        'total_amount': total_amount
    })


# @login_required
# def mobile_money_checkout(request):
#     # Vérifier si le panier est vide
#     cart = get_or_create_cart(request.user)
#     if cart.get_item_count() == 0:
#         return redirect('core/cart_detail')  # Si le panier est vide, on redirige

#     # Calculer le total des articles et le montant total
#     total_items = cart.get_item_count()
#     total_amount = sum(item.product.price_with_commission * item.quantity for item in cart.items.all())

#     if request.method == 'POST':
#         form = MobileMoneyPaymentForm(request.POST)
#         if form.is_valid():
#             # Sauvegarder le paiement Mobile Money
#             payment = form.save(commit=False)
#             payment.user = request.user
#             payment.status = 'pending'  # Statut initial en attente
#             payment.save()

#             # Créer l'ordre mais avec 'activated' = False (désactivé pour Mobile Money)
#             order = Order.objects.create(
#                 user=request.user,
#                 store=cart.items.first().product.store,  # Le store est celui du premier produit du panier
#                 status='pending',
#                 activated=False  # Désactivé par défaut pour Mobile Money
#             )

#             # Ajouter les articles de la commande
#             for cart_item in cart.items.all():
#                 OrderItem.objects.create(
#                     order=order,
#                     product=cart_item.product,
#                     quantity=cart_item.quantity,
#                     price_at_time_of_order=cart_item.product.price_with_commission
#                 )

#             # Calculer le total de la commande
#             order.calculate_total()

#             # Vider le panier de l'utilisateur après la commande
#             cart.items.all().delete()  # Suppression des articles dans le panier

#             # Message de succès et redirection vers la page d'attente
#             messages.success(request, 'Votre paiement Mobile Money a été soumis avec succès. Veuillez attendre la validation de l\'admin.')
#             return redirect('mobile_money_waiting', payment_id=payment.id)
#     else:
#         form = MobileMoneyPaymentForm()

#     return render(request, 'core/mobile_money_checkout.html', {
#         'form': form,
#         'total_items': total_items,
#         'total_amount': total_amount
#     })



import stripe
from django.conf import settings
from django.shortcuts import render, redirect
from django.http import JsonResponse
from .models import Cart
from django.contrib.auth.decorators import login_required

# Configurer Stripe avec la clé secrète
stripe.api_key = settings.STRIPE_TEST_SECRET_KEY

@login_required
def checkout(request):
    # Récupérer le panier de l'utilisateur
    cart = get_or_create_cart(request.user)
    if cart.get_item_count() == 0:
        return redirect('core/cart_detail')  # Si le panier est vide, on redirige
    
    # Taux de conversion CDF -> USD
    cdf_to_usd_rate = 2800
    # Créer une session de checkout de Stripe
    checkout_session = stripe.checkout.Session.create(
        payment_method_types=['card'],
        line_items=[
            {
                'price_data': {
                    'currency': 'eur',  # Monnaie de la transaction
                    'product_data': {
                        'name': item.product.name,
                    },
                    'unit_amount': int((item.product.price_with_commission * 100) / cdf_to_usd_rate),  # Montant en cents (USD)  # Montant en cents
                },
                'quantity': item.quantity,
            }
            for item in cart.items.all()
        ],
        mode='payment',
        success_url=request.build_absolute_uri('http://127.0.0.1:8000/user/payment_success/'),
        cancel_url=request.build_absolute_uri('http://127.0.0.1:8000/user/payment_cancel/'),
    )

    # Rediriger l'utilisateur vers la page de paiement Stripe
    return redirect(checkout_session.url, code=303)

# Page de succès après un paiement réussi


# Page de cancellation en cas d'annulation de paiement
def payment_cancel(request):
    return render(request, 'core/payment_cancel.html')
# @login_required
# def checkout(request):
#     cart = get_or_create_cart(request.user)
#     if cart.get_item_count() == 0:
#         return redirect('shop:cart_detail')  # Rediriger si le panier est vide
   
#     # Ici, tu peux ajouter la logique de paiement (par exemple, intégration avec Stripe)
#     # Mais pour l'instant, nous simulons juste le processus de commande.

#     # Marquer le panier comme "commandé" pour éviter qu'il soit modifié après
#     cart.is_ordered = True
#     cart.save()
   
#     return render(request, 'shop/checkout.html', {'cart': cart})

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render, get_object_or_404
from .models import Store, Category, Product
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import Store, Product, Category

from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import Store, Product, Category

# def store_detail(request, slug):
#     # Récupère le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug)
    
#     # Récupère les catégories associées au store
#     categories = Category.objects.filter(store=store)

#     # Récupère les produits associés au store
#     products = Product.objects.filter(store=store)

#     # Filtre par catégorie
#     category_filter = request.GET.get('categorie', '')  # Remarquez que 'categorie' est la clé
#     if category_filter:
#         # Filtrer les produits par l'ID de la catégorie sélectionnée
#         products = products.filter(category__id=category_filter)  # Filtrage correct par nom de catégorie

#     # Filtre par nom de produit
#     product_name = request.GET.get('nom', '')  # 'nom' est la clé utilisée dans le template
#     if product_name:
#         products = products.filter(name__icontains=product_name)  # Filtrage par nom de produit

#     # Filtre par prix
#     prix_min = request.GET.get('prix_min', '')  # 'prix_min' récupéré de la requête
#     prix_max = request.GET.get('prix_max', '')  # 'prix_max' récupéré de la requête
#     if prix_min:
#         try:
#             prix_min = float(prix_min)  # S'assurer que c'est un nombre valide
#             products = products.filter(price__gte=prix_min)
#         except ValueError:
#             pass  # Ignorer si le prix min n'est pas valide
#     if prix_max:
#         try:
#             prix_max = float(prix_max)  # S'assurer que c'est un nombre valide
#             products = products.filter(price__lte=prix_max)
#         except ValueError:
#             pass  # Ignorer si le prix max n'est pas valide

#     # Pagination
#     paginator = Paginator(products, 6)  # 6 produits par page
#     page = request.GET.get('page')

#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Passer les données au template
#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#     }

#     return render(request, 'core/store_detail.html', context)

from django.core.paginator import Paginator
from django.shortcuts import render, get_object_or_404, redirect
from .models import Store, Category, Product, Testimonial
from django.db.models import Count, Sum
from .forms import TestimonialForm
from django.db.models.functions import TruncDate
from django.db.models import Avg
from django.db.models import QuerySet

from django.shortcuts import get_object_or_404, render, redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Count, Sum
from datetime import datetime
from django.contrib import messages
from .models import Store, Category, Product, Order, Testimonialproduct, Testimonial
from .forms import TestimonialForm

from django.db.models import Sum, Count, F
from django.db.models.functions import TruncDate
from django.shortcuts import render, get_object_or_404
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from datetime import datetime
from django.contrib import messages
from .models import Store, Category, Product, Testimonialproduct, Testimonial, Order, OrderItem
from django.utils import timezone
from .models import Store, StoreVisit
from .utils import get_client_ip
from datetime import timedelta

def store_detail(request, slug):
    # Récupérer le store en fonction du slug
    store = get_object_or_404(Store, slug=slug)
    
    ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    # Récupérer les catégories associées au store
    categories = Category.objects.filter(store=store)

    # Récupérer les produits associés au store
    products = Product.objects.filter(store=store).order_by('-created_at')
    featured_products = products
    # Calculer la moyenne des notes pour chaque produit
    for product in products:
        testimonials = Testimonialproduct.objects.filter(product=product)
        if testimonials.exists():
            product.average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
        else:
            product.average_rating = 0

    # Créer une plage de 1 à 10 pour les étoiles
    range_10 = range(1, 11)

    # Filtre par catégorie
    category_filter = request.GET.get('categorie', '')
    if category_filter:
        products = products.filter(category__id=category_filter)

    # Filtre par nom de produit
    product_name = request.GET.get('nom', '').strip()
    if product_name:
        product_name = ' '.join(product_name.split())  # Normaliser l'espace
        products = products.filter(name__icontains=product_name) | products.filter(description__icontains=product_name)

    # Filtre par prix
    prix_min = request.GET.get('prix_min', '')
    prix_max = request.GET.get('prix_max', '')
    if prix_min:
        try:
            prix_min = float(prix_min)
            products = products.filter(price__gte=prix_min)
        except ValueError:
            pass
    if prix_max:
        try:
            prix_max = float(prix_max)
            products = products.filter(price__lte=prix_max)
        except ValueError:
            pass

    # Pagination pour les produits
    paginator = Paginator(products, 6)
    page = request.GET.get('page')
    try:
        products = paginator.page(page)
    except PageNotAnInteger:
        products = paginator.page(1)
    except EmptyPage:
        products = paginator.page(paginator.num_pages)

    # Formulaire de témoignage
    if request.method == 'POST':
        form = TestimonialForm(request.POST)
        if form.is_valid():
            testimonial = form.save(commit=False)
            testimonial.store = store
            testimonial.user = request.user
            testimonial.save()
            return redirect('store_detail', slug=slug)  # Redirige après ajout du témoignage
    else:
        form = TestimonialForm()
    
     # Calcul de la moyenne

    # Pour afficher les étoiles
    
    # Pagination des témoignages
    testimonials = Testimonial.objects.filter(store=store)
    average_rating = testimonials.aggregate(Avg('rating'))['rating__avg'] or 0
    rounded_rating = round(average_rating)
    testimonial_paginator = Paginator(testimonials, 3)
    testimonial_page = request.GET.get('testimonial_page')
    try:
        testimonials = testimonial_paginator.page(testimonial_page)
    except PageNotAnInteger:
        testimonials = testimonial_paginator.page(1)
    except EmptyPage:
        testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)

    # Initialiser les commandes par date, filtrées par les produits du store
    orders_by_date = Order.objects.filter(
        items__product__store=store,  # Assurez-vous d'utiliser 'items' si c'est le nom du champ qui lie Order à OrderItem
        activated=True
    ).annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only

    # Si la barre de recherche est utilisée pour filtrer par date
    order_date = request.GET.get('order_date', None)
    if order_date:
        try:
            # Convertir la date en format valide
            order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
            orders_by_date = orders_by_date.filter(order_date=order_date)
        except ValueError:
            messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
            orders_by_date = []

    # Appliquer l'agrégation pour compter les commandes et calculer le montant total
    # Agrégation pour compter les commandes distinctes et calculer le montant total
    # Agrégation pour compter les commandes distinctes et calculer le montant total
    orders_by_date = orders_by_date.values('order_date').annotate(
    total_orders=Count('id', distinct=True),  # Compte les commandes distinctes
    total_amount_sum=Sum(
        F('items__price_at_time_of_order') * F('items__quantity')
    )  # Calculer le montant total des produits pour chaque commande
    ).order_by('-order_date')


    # Pagination des commandes par date
    order_paginator = Paginator(orders_by_date, 6)
    order_page = request.GET.get('order_page')
    try:
        orders_by_date_page = order_paginator.page(order_page)
    except PageNotAnInteger:
        orders_by_date_page = order_paginator.page(1)
    except EmptyPage:
        orders_by_date_page = order_paginator.page(order_paginator.num_pages)

    # Vérification de la méthode count sur des QuerySets (pour éviter TypeError)
    product_count = products.count() if isinstance(products, QuerySet) else len(products)
    category_count = categories.count() if isinstance(categories, QuerySet) else len(categories)
    # Récupérer tous les produits avant pagination
    # Nombre total de produits sans pagination
    total_products = Product.objects.filter(store=store).count()

    today = timezone.now().date()
    week_ago = today - timedelta(days=6)  # 7 jours avec aujourd'hui

    daily_visits = StoreVisit.objects.filter(store=store, date=today).count()
    weekly_visits = StoreVisit.objects.filter(store=store, date__gte=week_ago).count()

    if request.user.is_authenticated:
        visit, created = StoreVisit.objects.get_or_create(
            store=store,
            user=request.user,
            date=today,
            defaults={'count': 1}
        )
    else:
        ip_address = get_client_ip(request)
        visit, created = StoreVisit.objects.get_or_create(
            store=store,
            user=None,
            ip_address=ip_address,
            date=today,
            defaults={'count': 1}
        )

# Récupérer tous les produits avant pagination
     # Calcul de la note moyenne de chaque store à l'aide d'agrégation
    

    # Passer tous les contextes nécessaires à la vue
    context = {
        'store': store,
        'categories': categories,
        'products': products,
        'paginator': paginator,
        'category_filter': category_filter,
        'product_name': product_name,
        'prix_min': prix_min,
        'prix_max': prix_max,
        'form': form,
        'testimonials': testimonials,
        'testimonial_paginator': testimonial_paginator,
        'range_10': range_10,
        'orders_by_date': orders_by_date_page,
        'order_date': order_date, 
        'order_paginator': order_paginator,
        'product_count': product_count,
        'total_products': total_products,
        'category_count': category_count,
        'daily_visits': daily_visits,
        'weekly_visits': weekly_visits,
        'favorite_stores':favorite_stores,
        'ad_popup': ad_popup,
        'average_rating': average_rating,
        'rounded_rating': rounded_rating,
        'rating_choices': Testimonial.RATING_CHOICES,
        'range_10': range(1, 11),
        'average_rating': average_rating,
        'rounded_rating': rounded_rating,
        'range_10': range(1, 11),
        'range_10': range(1, 11),
        'featured_products': featured_products
    }

    return render(request, 'core/store_detail.html', context)

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Store, SpotPubStore
from .forms import SpotPubStoreForm

def add_or_update_spotpub(request, slug):
    store = get_object_or_404(Store, slug=slug)

    try:
        spotpub = store.spot_pub  # OneToOneField
    except SpotPubStore.DoesNotExist:
        spotpub = None

    if request.method == 'POST':
        form = SpotPubStoreForm(request.POST, request.FILES, instance=spotpub)
        if form.is_valid():
            spot_instance = form.save(commit=False)
            spot_instance.store = store
            spot_instance.save()
            messages.success(request, "Vidéo publicitaire enregistrée avec succès.")
            return redirect('store_detail', slug=slug)
        else:
            messages.error(request, "La vidéo dépasse la taille maximale autorisée (10 Mo)..")
    else:
        form = SpotPubStoreForm(instance=spotpub)

    return render(request, 'core/add_or_update_spotpub.html', {
        'form': form,
        'store': store,
        'spotpub': spotpub,
        'video_url': spotpub.video.url if spotpub and spotpub.video else None,
    })





# def store_detail(request, slug):
#     # Récupérer le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug)

#     # Récupérer les catégories associées au store
#     categories = Category.objects.filter(store=store)

#     # Récupérer les produits associés au store
#     products = Product.objects.filter(store=store).order_by('-created_at')

#     # Calculer la moyenne des notes pour chaque produit
#     for product in products:
#         testimonials = Testimonialproduct.objects.filter(product=product)
#         if testimonials.exists():
#             product.average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
#         else:
#             product.average_rating = 0

#     # Créer une plage de 1 à 10 pour les étoiles
#     range_10 = range(1, 11)

#     # Filtre par catégorie
#     category_filter = request.GET.get('categorie', '')
#     if category_filter:
#         products = products.filter(category__id=category_filter)

#     # Filtre par nom de produit
#     product_name = request.GET.get('nom', '').strip()
#     if product_name:
#         product_name = ' '.join(product_name.split())  # Normaliser l'espace
#         products = products.filter(name__icontains=product_name) | products.filter(description__icontains=product_name)

#     # Filtre par prix
#     prix_min = request.GET.get('prix_min', '')
#     prix_max = request.GET.get('prix_max', '')
#     if prix_min:
#         try:
#             prix_min = float(prix_min)
#             products = products.filter(price__gte=prix_min)
#         except ValueError:
#             pass
#     if prix_max:
#         try:
#             prix_max = float(prix_max)
#             products = products.filter(price__lte=prix_max)
#         except ValueError:
#             pass

#     # Pagination pour les produits
#     paginator = Paginator(products, 6)
#     page = request.GET.get('page')
#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Formulaire de témoignage
#     if request.method == 'POST':
#         form = TestimonialForm(request.POST)
#         if form.is_valid():
#             testimonial = form.save(commit=False)
#             testimonial.store = store
#             testimonial.user = request.user
#             testimonial.save()
#             return redirect('store_detail', slug=slug)  # Redirige après ajout du témoignage
#     else:
#         form = TestimonialForm()

#     # Pagination des témoignages
#     testimonials = Testimonial.objects.filter(store=store)
#     testimonial_paginator = Paginator(testimonials, 3)
#     testimonial_page = request.GET.get('testimonial_page')
#     try:
#         testimonials = testimonial_paginator.page(testimonial_page)
#     except PageNotAnInteger:
#         testimonials = testimonial_paginator.page(1)
#     except EmptyPage:
#         testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)

#     # Initialiser les commandes par date, filtrées par les produits du store
#     orders_by_date = Order.objects.filter(
#         items__product__store=store,  # Assurez-vous d'utiliser 'items' si c'est le nom du champ qui lie Order à OrderItem
#         activated=True
#     ).annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only

#     # Si la barre de recherche est utilisée pour filtrer par date
#     order_date = request.GET.get('order_date', None)
#     if order_date:
#         try:
#             # Convertir la date en format valide
#             order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
#             orders_by_date = orders_by_date.filter(order_date=order_date)
#         except ValueError:
#             messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
#             orders_by_date = []

#     # Appliquer l'agrégation pour compter les commandes et calculer le montant total
#     orders_by_date = orders_by_date.values('order_date').annotate(
#         total_orders=Count('id'),
#         total_amount_sum=Sum('total_amount')
#     ).order_by('-order_date')

#     # Pagination des commandes par date
#     order_paginator = Paginator(orders_by_date, 5)
#     order_page = request.GET.get('order_page')
#     try:
#         orders_by_date_page = order_paginator.page(order_page)
#     except PageNotAnInteger:
#         orders_by_date_page = order_paginator.page(1)
#     except EmptyPage:
#         orders_by_date_page = order_paginator.page(order_paginator.num_pages)

#     # Vérification de la méthode count sur des QuerySets (pour éviter TypeError)
#     product_count = products.count() if isinstance(products, QuerySet) else len(products)
#     category_count = categories.count() if isinstance(categories, QuerySet) else len(categories)

#     # Passer tous les contextes nécessaires à la vue
#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#         'form': form,
#         'testimonials': testimonials,
#         'testimonial_paginator': testimonial_paginator,
#         'range_10': range_10,
#         'orders_by_date': orders_by_date_page,
#         'order_date': order_date, 
#         'order_paginator': order_paginator,
#         'product_count': product_count,
#         'category_count': category_count,
#     }

#     return render(request, 'core/store_detail.html', context)

# def store_detail(request, slug):
#     # Récupérer le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug)

#     # Récupérer les catégories associées au store
#     categories = Category.objects.filter(store=store)

#     # Récupérer les produits associés au store
#     products = Product.objects.filter(store=store).order_by('-created_at')

#     # Calculer la moyenne des notes pour chaque produit
#     for product in products:
#         testimonials = Testimonialproduct.objects.filter(product=product)
#         if testimonials.exists():
#             product.average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
#         else:
#             product.average_rating = 0

#     # Créer une plage de 1 à 10 pour les étoiles
#     range_10 = range(1, 11)

#     # Filtre par catégorie
#     category_filter = request.GET.get('categorie', '')
#     if category_filter:
#         products = products.filter(category__id=category_filter)

#     # Filtre par nom de produit
#     product_name = request.GET.get('nom', '').strip()
#     if product_name:
#         product_name = ' '.join(product_name.split())  # Normaliser l'espace
#         products = products.filter(name__icontains=product_name) | products.filter(description__icontains=product_name)

#     # Filtre par prix
#     prix_min = request.GET.get('prix_min', '')
#     prix_max = request.GET.get('prix_max', '')
#     if prix_min:
#         try:
#             prix_min = float(prix_min)
#             products = products.filter(price__gte=prix_min)
#         except ValueError:
#             pass
#     if prix_max:
#         try:
#             prix_max = float(prix_max)
#             products = products.filter(price__lte=prix_max)
#         except ValueError:
#             pass

#     # Pagination pour les produits
#     paginator = Paginator(products, 6)
#     page = request.GET.get('page')
#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Formulaire de témoignage
#     if request.method == 'POST':
#         form = TestimonialForm(request.POST)
#         if form.is_valid():
#             testimonial = form.save(commit=False)
#             testimonial.store = store
#             testimonial.user = request.user
#             testimonial.save()
#             return redirect('store_detail', slug=slug)  # Redirige après ajout du témoignage
#     else:
#         form = TestimonialForm()

#     # Pagination des témoignages
#     testimonials = Testimonial.objects.filter(store=store)
#     testimonial_paginator = Paginator(testimonials, 3)
#     testimonial_page = request.GET.get('testimonial_page')
#     try:
#         testimonials = testimonial_paginator.page(testimonial_page)
#     except PageNotAnInteger:
#         testimonials = testimonial_paginator.page(1)
#     except EmptyPage:
#         testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)

#     # Initialiser les commandes par date, filtrées par les produits du store
#     orders_by_date = Order.objects.filter(
#         orderitem__product__store=store,  # Filtrer les commandes qui contiennent des produits du store
#         activated=True
#     ).annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only

#     # Si la barre de recherche est utilisée pour filtrer par date
#     order_date = request.GET.get('order_date', None)
#     if order_date:
#         try:
#             # Convertir la date en format valide
#             order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
#             orders_by_date = orders_by_date.filter(order_date=order_date)
#         except ValueError:
#             messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
#             orders_by_date = []

#     # Appliquer l'agrégation pour compter les commandes et calculer le montant total
#     orders_by_date = orders_by_date.values('order_date').annotate(
#         total_orders=Count('id'),
#         total_amount_sum=Sum('total_amount')
#     ).order_by('-order_date')

#     # Pagination des commandes par date
#     order_paginator = Paginator(orders_by_date, 5)
#     order_page = request.GET.get('order_page')
#     try:
#         orders_by_date_page = order_paginator.page(order_page)
#     except PageNotAnInteger:
#         orders_by_date_page = order_paginator.page(1)
#     except EmptyPage:
#         orders_by_date_page = order_paginator.page(order_paginator.num_pages)

#     # Passer tous les contextes nécessaires à la vue
#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#         'form': form,
#         'testimonials': testimonials,
#         'testimonial_paginator': testimonial_paginator,
#         'range_10': range_10,
#         'orders_by_date': orders_by_date_page,
#         'order_date': order_date, 
#         'order_paginator': order_paginator,
#         'product_count': products.count(),
#         'category_count': categories.count(),
#     }

#     return render(request, 'core/store_detail.html', context)


# def store_detail(request, slug):
#     # Récupère le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug)
    
#     # Récupère les catégories associées au store
#     categories = Category.objects.filter(store=store)

#     # Récupère les produits associés au store
#     products = Product.objects.filter(store=store).order_by('-created_at')
#     for product in products:
#         # Récupère tous les témoignages pour ce produit
#         testimonials = Testimonialproduct.objects.filter(product=product)
        
#         # Calcule la moyenne des notes si des témoignages existent
#         if testimonials.exists():
#             product.average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
#         else:
#             product.average_rating = 0
#     # Créez une plage de 1 à 10 pour les étoiles
#     range_10 = range(1, 11)
    
#     # Filtre par catégorie
#     category_filter = request.GET.get('categorie', '')
#     if category_filter:
#         products = products.filter(category__id=category_filter)

#     # Filtre par nom de produit
#     product_name = request.GET.get('nom', '').strip()  # On enlève les espaces en début et fin
#     if product_name:
#     # On remplace les espaces multiples par un espace unique et on effectue la recherche
#        product_name = ' '.join(product_name.split())
#        products = products.filter(name__icontains=product_name) | products.filter(description__icontains=product_name)

#     # product_name = request.GET.get('nom', '')
#     # if product_name:
#     #     products = products.filter(name__icontains=product_name)

#     # Filtre par prix
#     prix_min = request.GET.get('prix_min', '')
#     prix_max = request.GET.get('prix_max', '')
#     if prix_min:
#         try:
#             prix_min = float(prix_min)
#             products = products.filter(price__gte=prix_min)
#         except ValueError:
#             pass
#     if prix_max:
#         try:
#             prix_max = float(prix_max)
#             products = products.filter(price__lte=prix_max)
#         except ValueError:
#             pass

#     # Pagination pour les produits
#     paginator = Paginator(products, 6)
#     page = request.GET.get('page')
#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Formulaire de témoignage
#     if request.method == 'POST':
#         form = TestimonialForm(request.POST)
#         if form.is_valid():
#             testimonial = form.save(commit=False)
#             testimonial.store = store
#             testimonial.user = request.user
#             testimonial.save()
#             return redirect('store_detail', slug=slug)  # Redirige après ajout du témoignage
#     else:
#         form = TestimonialForm()

#     # Pagination des témoignages
#     testimonials = Testimonial.objects.filter(store=store)
#     testimonial_paginator = Paginator(testimonials, 3)
#     testimonial_page = request.GET.get('testimonial_page')
#     try:
#         testimonials = testimonial_paginator.page(testimonial_page)
#     except PageNotAnInteger:
#         testimonials = testimonial_paginator.page(1)
#     except EmptyPage:
#         testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)
    
#     # Initialiser les commandes
#     orders_by_date = Order.objects.filter(store=store, activated=True) \
#                                   .annotate(order_date=TruncDate('created_at'))  # Truncate 'created_at' to date only
    
#     # Si la barre de recherche est utilisée pour filtrer par date
#     order_date = request.GET.get('order_date', None)
#     if order_date:
#         try:
#             # Convertir la date en format valide
#             order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
#             orders_by_date = orders_by_date.filter(order_date=order_date)  # Filtrer par la date choisie
#         except ValueError:
#             messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
#             orders_by_date = []

#     # Appliquer l'agrégation pour compter les commandes et calculer le montant total
#     orders_by_date = orders_by_date.values('order_date')  # Regroupement par date
#     orders_by_date = orders_by_date.annotate(
#         total_orders=Count('id'),  # Nombre de commandes pour cette date
#         total_amount_sum=Sum('total_amount')   # Somme des montants des commandes
#     ).order_by('-order_date')  # Tri décroissant par date

#     # Pagination des commandes par date
#     order_paginator = Paginator(orders_by_date, 5)
#     order_page = request.GET.get('order_page')
#     try:
#         orders_by_date_page = order_paginator.page(order_page)
#     except PageNotAnInteger:
#         orders_by_date_page = order_paginator.page(1)
#     except EmptyPage:
#         orders_by_date_page = order_paginator.page(order_paginator.num_pages)

#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#         'form': form,
#         'testimonials': testimonials,
#         'testimonial_paginator': testimonial_paginator,
#         'range_10': range_10,
#         'orders_by_date': orders_by_date_page,
#         'order_date': order_date, 
#         'order_paginator': order_paginator,
#     }

#     return render(request, 'core/store_detail.html', context)

# manage product
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Store, Product, Category
from .forms import ProductForm
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
@login_required
def manage_product_store(request, slug):
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    products = Product.objects.filter(store=store).order_by('-created_at')

    categories = Category.objects.filter(store=store)

    category_filter = request.GET.get('categorie', '')
    if category_filter:
        products = products.filter(category__id=category_filter)

    product_name = request.GET.get('nom', '')
    if product_name:
        products = products.filter(name__icontains=product_name)

    paginator = Paginator(products, 6)
    page = request.GET.get('page')
    try:
        products = paginator.page(page)
    except PageNotAnInteger:
        products = paginator.page(1)
    except EmptyPage:
        products = paginator.page(paginator.num_pages)

    context = {
        'store': store,
        'products': products,
        'categories': categories,
        'paginator': paginator,
    }

    return render(request, 'core/manage_product_store.html', context)

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Product
from .forms import ProductForm


from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Product, Photo
from .forms import ProductForm, PhotoForm
@login_required

@login_required
def edit_product(request, product_id):
    # Récupère le produit à partir de son ID
    product = get_object_or_404(Product, id=product_id)

    # Vérifie que l'utilisateur est le propriétaire du magasin
    if product.store.owner != request.user:
        messages.error(request, "Vous n'êtes pas autorisé à modifier ce produit.")
        return redirect('manage_product_store', slug=product.store.slug)

    # Formulaire pour modifier les informations du produit
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES, instance=product)
        if form.is_valid():
            form.save()
            messages.success(request, "Produit mis à jour avec succès.")
            
            # Traitement des images pour la galerie
            if 'image_galerie' in request.FILES:
                for image in request.FILES.getlist('image_galerie'):
                    photo = Photo(product=product, image=image)
                    photo.save()
                messages.success(request, "Images ajoutées à la galerie avec succès.")

            # Rediriger vers la page de gestion des produits du magasin
            return redirect('manage_product_store', slug=product.store.slug)
        else:
            messages.error(request, "Il y a des erreurs dans le formulaire.")
    else:
        form = ProductForm(instance=product)

    # Récupère les photos existantes associées au produit
    photos = product.photos.all()

    # Formulaire pour ajouter des photos à la galerie
    photo_form = PhotoForm()

    context = {
        'form': form,
        'photo_form': photo_form,
        'product': product,
        'photos': photos,
    }

    return render(request, 'core/edit_product.html', context)



from django.shortcuts import get_object_or_404, redirect
from .models import Photo

def delete_photo(request, photo_id):
    photo = get_object_or_404(Photo, id=photo_id)
    product = photo.product
    if request.method == 'POST':
        photo.delete()
        return redirect('edit_product', product_id=product.id)


from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Product

@login_required
def delete_product(request, product_id):
    product = get_object_or_404(Product, id=product_id)

    # Vérifie que l'utilisateur est le propriétaire du produit ou du magasin
    if product.store.owner != request.user:
        messages.error(request, "Vous n'êtes pas autorisé à supprimer ce produit.")
        return redirect('manage_product_store', slug=product.store.slug)

    if request.method == 'POST':
        product.delete()
        messages.success(request, "Produit supprimé avec succès.")
        return redirect('manage_product_store', slug=product.store.slug)
    else:
        messages.error(request, "La suppression a échoué. Essayez à nouveau.")
        return redirect('manage_product_store', slug=product.store.slug)


# @login_required
# def store_detail(request, slug):
#     # Récupère le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug)
    
#     # Récupère les catégories associées au store
#     categories = Category.objects.filter(store=store)

#     # Récupère les produits associés au store
#     products = Product.objects.filter(store=store)
#     for product in products:
#         # Récupère tous les témoignages pour ce produit
#         testimonials = Testimonialproduct.objects.filter(product=product)
        
#         # Calcule la moyenne des notes si des témoignages existent
#         if testimonials.exists():
#             product.average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
#         else:
#             product.average_rating = 0
#     # Créez une plage de 1 à 10 pour les étoiles
#     range_10 = range(1, 11)
    
#     # Filtre par catégorie
#     category_filter = request.GET.get('categorie', '')
#     if category_filter:
#         products = products.filter(category__id=category_filter)

#     # Filtre par nom de produit
#     product_name = request.GET.get('nom', '')
#     if product_name:
#         products = products.filter(name__icontains=product_name)

#     # Filtre par prix
#     prix_min = request.GET.get('prix_min', '')
#     prix_max = request.GET.get('prix_max', '')
#     if prix_min:
#         try:
#             prix_min = float(prix_min)
#             products = products.filter(price__gte=prix_min)
#         except ValueError:
#             pass
#     if prix_max:
#         try:
#             prix_max = float(prix_max)
#             products = products.filter(price__lte=prix_max)
#         except ValueError:
#             pass

#     # Pagination pour les produits
#     paginator = Paginator(products, 6)
#     page = request.GET.get('page')
#     try:
#         products = paginator.page(page)
#     except PageNotAnInteger:
#         products = paginator.page(1)
#     except EmptyPage:
#         products = paginator.page(paginator.num_pages)

#     # Formulaire de témoignage
#     if request.method == 'POST':
#         form = TestimonialForm(request.POST)
#         if form.is_valid():
#             testimonial = form.save(commit=False)
#             testimonial.store = store
#             testimonial.user = request.user
#             testimonial.save()
#             return redirect('store_detail', slug=slug)  # Redirige après ajout du témoignage
#     else:
#         form = TestimonialForm()

#     # Pagination des témoignages
#     testimonials = Testimonial.objects.filter(store=store)
#     testimonial_paginator = Paginator(testimonials, 3)
#     testimonial_page = request.GET.get('testimonial_page')
#     try:
#         testimonials = testimonial_paginator.page(testimonial_page)
#     except PageNotAnInteger:
#         testimonials = testimonial_paginator.page(1)
#     except EmptyPage:
#         testimonials = testimonial_paginator.page(testimonial_paginator.num_pages)
    
#     # initialiser
#     # Initialiser les commandes avec toutes les commandes pour ce magasin

#     # Récupérer la date spécifiée dans la barre de recherche (format yyyy-mm-dd)
    

#     # Si la barre de recherche est utilisée pour filtrer par date
#     orders_by_date = Order.objects.filter(store=store) \
#                                   .annotate(order_date=TruncDate('created_at')) \
#                                   .values('order_date') \
#                                   .distinct()

#     # Si la barre de recherche est utilisée pour filtrer par date
#     order_date = request.GET.get('order_date', None)
#     if order_date:
#         # Convertir la date en format valide
#         try:
#             order_date = datetime.strptime(order_date, '%Y-%m-%d').date()
#             orders_by_date = orders_by_date.filter(order_date=order_date)
#         except ValueError:
#             # Gérer le cas où la date est invalide
#             messages.error(request, "La date fournie est invalide. Veuillez entrer une date correcte.")
#             orders_by_date = []

#     # order par date
    
#     # Grouper les commandes par date (sans l'heure)
#     orders = Order.objects.filter(store=store)

#     # Regroupe par date en utilisant TruncDate et effectue les agrégations
#     orders_by_date = orders.annotate(
#         order_date=TruncDate('created_at')  # Truncate the timestamp to date
#     ).values('order_date')  # Extract the date for grouping

#     # Applique l'agrégation pour compter le nombre de commandes et la somme du montant
#     orders_by_date = orders_by_date.annotate(
#         total_orders=Count('id'),
#         total_amount=Sum('total_amount')
#     ).order_by('-order_date')  # Tri décroissant par date


#     # Regroupe par date en utilisant TruncDate et effectue les agrégations
#     orders_by_date = orders.annotate(
#         order_date=TruncDate('created_at')  # Truncate the timestamp to date
#     ).values('order_date')  # Extract the date for grouping

#     # Applique l'agrégation pour compter le nombre de commandes et la somme du montant
#     orders_by_date = orders_by_date.annotate(
#         total_orders=Count('id'),
#         total_amount=Sum('total_amount')
#     ).order_by('-order_date')  # Tri décroissant par date

#     # Pagination pour afficher 5 dates par page
#     order_paginator = Paginator(orders_by_date, 5)
#     order_page = request.GET.get('order_page')
#     try:
#         orders_by_date_page = order_paginator.page(order_page)
#     except PageNotAnInteger:
#         orders_by_date_page = order_paginator.page(1)
#     except EmptyPage:
#         orders_by_date_page = order_paginator.page(order_paginator.num_pages)

#     context = {
#         'store': store,
#         'categories': categories,
#         'products': products,
#         'paginator': paginator,
#         'category_filter': category_filter,
#         'product_name': product_name,
#         'prix_min': prix_min,
#         'prix_max': prix_max,
#         'form': form,
#         'testimonials': testimonials,
#         'testimonial_paginator': testimonial_paginator,
#         'range_10': range_10,
#         'orders_by_date': orders_by_date_page,
#         'order_date': order_date, 
#         'order_paginator': order_paginator,
#     }

#     return render(request, 'core/store_detail.html', context)

from datetime import datetime
from django.shortcuts import render, get_object_or_404
from datetime import datetime
from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import Order, Store
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from datetime import datetime
from django.db.models import Sum, Count
from .models import Store, Order, OrderItem

@login_required
def orders_by_date_detail(request, slug, order_date):
    # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    # Convertir la date reçue dans l'URL en un objet date
    order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

    # Filtrer les OrderItems liés à ce store et à cette date
    order_items = OrderItem.objects.filter(
        product__store=store,  # Filtrer les OrderItems par le store du produit
        order__created_at__date=order_date,  # Filtrer par la date de création de la commande
        order__activated=True  # S'assurer que la commande est activée
    ).select_related('order', 'product')  # Sélectionner les objets associés pour optimisation

    # Organiser les données par client (utilisateur de la commande)
    orders_data = []
    total_orders = 0  # Total des commandes
    total_amount = 0  # Total montant des ventes global

    # Créer un dictionnaire pour suivre les commandes par client
    client_data = {}

    for order_item in order_items:
        client_name = order_item.order.user.username  # Nom du client
        total_items_in_order = order_item.quantity  # Nombre d'articles pour cet item de commande
        total_order_amount = order_item.price_at_time_of_order * order_item.quantity  # Montant pour cet item de commande

        # Ajouter ou mettre à jour les données du client
        if client_name not in client_data:
            client_data[client_name] = {
                'total_items': 0,  # Nombre total d'articles pour ce client
                'total_amount': 0  # Montant total pour ce client
            }

        client_data[client_name]['total_items'] += total_items_in_order
        client_data[client_name]['total_amount'] += total_order_amount

    # Remplir orders_data à partir de client_data
    for client_name, data in client_data.items():
        orders_data.append({
            'client_name': client_name,
            'total_items_in_order': data['total_items'],
            'total_order_amount': data['total_amount'],
        })

    # Calculer les totaux globaux
    total_orders = len(orders_data)  # Total des commandes
    total_amount = sum([data['total_order_amount'] for data in orders_data])  # Total montant des ventes

    # Pagination des commandes
    paginator = Paginator(orders_data, 6)  # 5 commandes par page
    page = request.GET.get('page')
    try:
        orders_data_page = paginator.page(page)
    except PageNotAnInteger:
        orders_data_page = paginator.page(1)
    except EmptyPage:
        orders_data_page = paginator.page(paginator.num_pages)

    # Passer les données à l'HTML pour l'affichage
    context = {
        'store': store,
        'orders_data': orders_data_page,
        'total_orders': total_orders,
        'total_amount': total_amount,
        'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
    }

    return render(request, 'core/orders_by_date_detail.html', context)

# @login_required
# def orders_by_date_detail(request, slug, order_date):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)

#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

#     # Filtrer les OrderItems liés à ce store et à cette date
#     order_items = OrderItem.objects.filter(
#         product__store=store,  # Filtrer les OrderItems par le store du produit
#         order__created_at__date=order_date,  # Filtrer par la date de création de la commande
#         order__activated=True  # S'assurer que la commande est activée
#     ).select_related('order', 'product')  # Sélectionner les objets associés pour optimisation

#     # Calculer le nombre total de commandes et le montant total
#     total_orders = order_items.count()
#     total_amount = sum(item.get_total_price() for item in order_items)

#     # Organiser les données par commande et récupérer les informations nécessaires
#     orders_data = []
#     for order in order_items.values('order').distinct():
#         order_instance = Order.objects.get(id=order['order'])
#         # Nombre total d'articles dans cette commande
#         total_items_in_order = order_instance.items.count()
#         # Montant total de la commande
#         total_order_amount = order_instance.total_amount
#         # Nom du client
#         client_name = order_instance.user.username

#         orders_data.append({
#             'client_name': client_name,
#             'total_items_in_order': total_items_in_order,
#             'total_order_amount': total_order_amount,
#         })

#     # Pagination des commandes
#     paginator = Paginator(orders_data, 5)  # 5 commandes par page
#     page = request.GET.get('page')
#     try:
#         orders_data_page = paginator.page(page)
#     except PageNotAnInteger:
#         orders_data_page = paginator.page(1)
#     except EmptyPage:
#         orders_data_page = paginator.page(paginator.num_pages)

#     # Passer les données à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'orders_data': orders_data_page,
#         'total_orders': total_orders,
#         'total_amount': total_amount,
#         'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
#     }

#     return render(request, 'core/orders_by_date_detail.html', context)

# @login_required
# def orders_by_date_detail(request, slug, order_date):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)

#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

#     # Filtrer les OrderItem liés à ce store et à cette date
#     order_items = OrderItem.objects.filter(
#         product__store=store,  # Filtrer les OrderItem par store via le produit associé
#         order__created_at__date=order_date,  # Filtrer par la date de création de la commande
#         order__activated=True  # Assurer que la commande est activée
#     ).select_related('order', 'product')  # Sélectionner les objets associés pour optimisation

#     # Calculer le nombre total de commandes et le montant total
#     total_orders = order_items.count()
#     total_amount = sum(item.get_total_price() for item in order_items)

#     # Pagination des articles de commande (OrderItems)
#     paginator = Paginator(order_items, 5)  # 5 OrderItems par page
#     page = request.GET.get('page')
#     try:
#         order_items_page = paginator.page(page)
#     except PageNotAnInteger:
#         order_items_page = paginator.page(1)
#     except EmptyPage:
#         order_items_page = paginator.page(paginator.num_pages)

#     # Passer la date à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'order_items': order_items_page,
#         'total_orders': total_orders,
#         'total_amount': total_amount,
#         'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
#     }

#     return render(request, 'core/orders_by_date_detail.html', context)


# @login_required
# def orders_by_date_detail(request, slug, order_date):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)

#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

#     # Filtrer les OrderItem liés à ce store et à cette date
#     order_items = OrderItem.objects.filter(
#         store=store, 
#         order__created_at__date=order_date,  # Filtrer les commandes par la date de création
#         order__activated=True
#     ).select_related('order')  # Nous sélectionnons 'order' pour avoir l'accès à l'objet complet

#     # Calculer le nombre total de commandes et le montant total
#     total_orders = order_items.count()
#     total_amount = sum(item.get_total_price() for item in order_items)

#     # Pagination des articles de commande
#     paginator = Paginator(order_items, 5)  # 5 articles de commande par page
#     page = request.GET.get('page')
#     try:
#         order_items_page = paginator.page(page)
#     except PageNotAnInteger:
#         order_items_page = paginator.page(1)
#     except EmptyPage:
#         order_items_page = paginator.page(paginator.num_pages)

#     # Passer la date à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'order_items': order_items_page,
#         'total_orders': total_orders,
#         'total_amount': total_amount,
#         'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
#     }

#     return render(request, 'core/orders_by_date_detail.html', context)


# @login_required
# def orders_by_date_detail(request, slug, order_date):
#     # Récupérer le store en fonction du slug
#     store = get_object_or_404(Store, slug=slug, owner=request.user)
    
#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()  # Conversion de la chaîne en objet date
    
#     # Filtrer les commandes par date et magasin
#     orders_queryset = Order.objects.filter(store=store, created_at__date=order_date, activated=True).order_by('-created_at')
    
#     # Calculer le nombre total de commandes et le montant total avant la pagination
#     total_orders = orders_queryset.count()  # Nombre total de commandes
#     total_amount = sum(order.total_amount for order in orders_queryset)  # Calcul du montant total

#     # Pagination des commandes
#     paginator = Paginator(orders_queryset, 5)  # Vous pouvez ajuster le nombre de commandes par page
#     page = request.GET.get('page')
#     try:
#         orders = paginator.page(page)
#     except PageNotAnInteger:
#         orders = paginator.page(1)
#     except EmptyPage:
#         orders = paginator.page(paginator.num_pages)

#     # Passer la date à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'orders': orders,
#         'total_orders': total_orders,
#         'total_amount': total_amount,
#         'order_date': order_date,  # Ajouter la date pour l'afficher dans le template
#     }

#     return render(request, 'core/orders_by_date_detail.html', context)
from django.shortcuts import render, get_object_or_404
from .models import Order
from django.contrib.auth import get_user_model
from .models import Order,CustomUser,ProductPoints

@login_required
def order_detail_store(request, slug, client_username, order_id):
    # Récupérer le store en fonction du slug et vérifier qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)

    # Récupérer l'utilisateur (client) en fonction du nom d'utilisateur
    User = get_user_model()
    user = get_object_or_404(User, username=client_username)

    # Récupérer la commande spécifique du client en fonction de l'ID de la commande et du store
    order = get_object_or_404(Order, id=order_id, user=user, items__product__store=store)

    # Récupérer les articles associés à cette commande
    order_items = order.items.filter(product__store=store)

    # Calculer le nombre d'articles et le montant total de la commande
    total_items_in_order = order_items.count()  # Nombre d'articles
    total_order_amount = order.total_amount  # Montant total de la commande

    # Passer les données à ton template
    context = {
        'store': store,
        'user': user,
        'order': order,
        'order_items': order_items,
        'total_items_in_order': total_items_in_order,
        'total_order_amount': total_order_amount,
    }

    return render(request, 'core/order_detail_store.html', context)
@login_required
def orders_by_date_detail_detail(request, slug, order_date, client_username):
    # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    user = get_object_or_404(CustomUser, username=client_username)

    # Convertir la date reçue dans l'URL en un objet date
    order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

    # Récupérer les commandes du client pour ce store et cette date
    orders = Order.objects.filter(
        user__username=client_username,
        created_at__date=order_date,
        items__product__store=store,
        activated=True
    ).distinct()

    # Créer une liste de détails de commande, avec les noms des articles
    order_data = []
    total_global_amount = 0  # Initialiser le montant total global

    for order in orders:
        # Récupérer tous les items associés à cette commande
        order_items = order.items.filter(product__store=store)

        # Initialiser les variables pour compter les articles et calculer le montant total
        total_items_in_order = 0
        total_order_amount = 0

        # Extraire les noms des articles et calculer le total des montants
        item_names = []
        for item in order_items:
            total_items_in_order += item.quantity  # Ajouter la quantité de cet article
            total_order_amount += item.price_at_time_of_order * item.quantity  # Ajouter le montant total pour cet article
            item_names.append(item.product.name)  # Ajouter le nom de l'article

        # Ajouter les informations nécessaires pour chaque commande
        order_data.append({
            'order_id': order.id,
            'total_items_in_order': total_items_in_order,
            'total_order_amount': total_order_amount,
            'item_names': ', '.join(item_names),  # Liste des noms des articles
        })

        # Ajouter au montant global
        total_global_amount += total_order_amount
        # Calculer le nombre total de commandes pour ce client
        total_orders = orders.count()

    # Passer les données à l'HTML pour l'affichage
    context = {
        'store': store,
        'order_data': order_data,
        'order_date': order_date,
        'client_username': client_username,
        'user_phone': user.phone,
        'total_global_amount': total_global_amount,  # Ajouter le montant total global au contexte
        'total_orders': total_orders  # Ajouter le nombre total de commandes
    }

    return render(request, 'core/orders_by_date_detail_detail.html', context)


from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import ProductPoints, UserPoints  # Assurez-vous d'importer UserPoints si vous avez une telle table

from django.contrib import messages
from django.shortcuts import render, redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import ProductPoints, UserPoints, Purchase
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import  UserPoints, Purchase
from django.http import JsonResponse

# def list_product_rewards(request):
#     # Récupère tous les produits de récompenses
#     product_rewards = ProductPoints.objects.all()

#     # Filtrage par nom et description
#     search_query = request.GET.get('search', '').strip()  # Recherche par nom et description
#     if search_query:
#         product_rewards = product_rewards.filter(name__icontains=search_query) | product_rewards.filter(description__icontains=search_query)

#     # Pagination
#     paginator = Paginator(product_rewards, 6)  # 6 produits par page
#     page = request.GET.get('page')
#     try:
#         product_rewards = paginator.page(page)
#     except PageNotAnInteger:
#         product_rewards = paginator.page(1)
#     except EmptyPage:
#         product_rewards = paginator.page(paginator.num_pages)

#     # Récupère les points de fidélité de l'utilisateur connecté
#     try:
#         user_points = request.user.userpoints.points  # Si vous utilisez une table 'UserPoints' pour les points
#     except UserPoints.DoesNotExist:
#         user_points = 0  # Si l'utilisateur n'a pas de points, on affiche 0

#     # Traitement de l'achat
#     if request.method == 'POST':
#         product_id = request.POST.get('product_id')
#         product = ProductPoints.objects.get(id=product_id)

#         # Vérifier si l'utilisateur a assez de points
#         if user_points >= product.points_required:
#             # Déduire les points de l'utilisateur
#             user_points_obj = UserPoints.objects.get(user=request.user)
#             user_points_obj.points -= product.points_required
#             user_points_obj.save()

#             # Enregistrer l'achat dans le modèle Purchase
#             purchase = Purchase.objects.create(
#                 user=request.user,
#                 product=product,
#                 points_used=product.points_required
#             )

#             # Ajouter un message de succès
#             messages.success(request, f"Achat de {product.name} effectué avec succès !")
#             return redirect('list_product_rewards')  # Vous pouvez rediriger vers la même page ou une page de confirmation

#         else:
#             messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
#             return redirect('list_product_rewards')  # Redirigez vers la même page pour afficher l'erreur

#     # Si l'utilisateur n'a pas encore tenté d'acheter
#     context = {
#         'product_rewards': product_rewards,
#         'search_query': search_query,
#         'user_points': user_points,  # Les points de l'utilisateur
#         'paginator': paginator,
#     }

#     return render(request, 'core/list_product_reward.html', context)
@login_required
def list_product_rewards(request):
    product_rewards = ProductPoints.objects.all().order_by('-created_at')
    featured_products = product_rewards
    ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    range_10 = range(1, 11)
    search_query = request.GET.get('search', '').strip()
    if search_query:
        product_rewards = product_rewards.filter(name__icontains=search_query) | product_rewards.filter(description__icontains=search_query)

    paginator = Paginator(product_rewards, 6)
    page = request.GET.get('page')
    try:
        product_rewards = paginator.page(page)
    except PageNotAnInteger:
        product_rewards = paginator.page(1)
    except EmptyPage:
        product_rewards = paginator.page(paginator.num_pages)

    try:
        user_points = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        user_points = None

    # Récupérer le taux de conversion
    try:
        conversion = PointConversion.objects.latest('id')  # Dernier taux de conversion ajouté
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5  # Valeur par défaut

    # Ajouter la valeur en USD à chaque produit
    for product in product_rewards:
        product.usd_price = product.points_required * conversion_rate  # Conversion

    if request.method == 'POST':
        product_id = request.POST.get('product_id')
        product = ProductPoints.objects.get(id=product_id)

        if user_points and user_points.spend_points(product.points_required):
            Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
            return redirect('list_product_rewards')
        else:
            messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")

    return render(request, 'core/list_product_reward.html', {
        'product_rewards': product_rewards,
        'search_query': search_query,
        'user_points': user_points.points if user_points else 0,
        'spent_points': user_points.spent_points if user_points else 0,
        'paginator': paginator,
        'favorite_stores':favorite_stores,
        'range_10': range_10, 
        'ad_popup': ad_popup,
        'featured_products': featured_products 
    })

from django.shortcuts import render, get_object_or_404
from .models import ProductPoints, PointConversion
@login_required
def detail_product_reward(request, product_id):
    product = get_object_or_404(ProductPoints, id=product_id)
    user_points = UserPoints.objects.get(user=request.user)

    try:
        conversion = PointConversion.objects.latest('id')
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5

    price_in_usd = product.points_required * conversion_rate

    if request.method == 'POST':
        # Vérifie si l'utilisateur a assez de points
        if user_points and user_points.points >= product.points_required:
            user_points.spend_points(product.points_required)
            Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
        else:
            messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
        
        # Redirige vers la liste quoi qu'il arrive (succès ou erreur)
        return redirect('list_product_rewards')

    return render(request, 'core/detail_product_reward.html', {
        'product': product,
        'user_points': user_points.points if user_points else 0,
        'price_in_usd': price_in_usd,
        'spent_points': user_points.spent_points if user_points else 0,
    })


# @login_required
# def detail_product_reward(request, product_id):
#     product = get_object_or_404(ProductPoints, id=product_id)
#      # Récupérer les photos associées
    
#     # Récupérer le taux de conversion des points
#     try:
#         conversion = PointConversion.objects.latest('id')  # Dernier taux défini
#         conversion_rate = conversion.conversion_rate
#     except PointConversion.DoesNotExist:
#         conversion_rate = 0.5  # Valeur par défaut

#     # Convertir les points en valeur monétaire
#     price_in_usd = product.points_required * conversion_rate

#     context = {
#         'product': product,
#         'price_in_usd': price_in_usd,
#     }
#     return render(request, 'core/detail_product_reward.html', context)  # ✅ Plus de virgule ici


# @login_required
# def list_product_rewards(request):
#     product_rewards = ProductPoints.objects.all()

#     search_query = request.GET.get('search', '').strip()
#     if search_query:
#         product_rewards = product_rewards.filter(name__icontains=search_query) | product_rewards.filter(description__icontains=search_query)

#     paginator = Paginator(product_rewards, 6)
#     page = request.GET.get('page')
#     try:
#         product_rewards = paginator.page(page)
#     except PageNotAnInteger:
#         product_rewards = paginator.page(1)
#     except EmptyPage:
#         product_rewards = paginator.page(paginator.num_pages)

#     try:
#         user_points = UserPoints.objects.get(user=request.user)
#     except UserPoints.DoesNotExist:
#         user_points = None

#     if request.method == 'POST':
#         product_id = request.POST.get('product_id')
#         product = ProductPoints.objects.get(id=product_id)

#         if user_points and user_points.spend_points(product.points_required):
#             Purchase.objects.create(
#                 user=request.user,
#                 product=product,
#                 points_used=product.points_required
#             )
#             messages.success(request, f"Achat de {product.name} effectué avec succès !")
#             return redirect('list_product_rewards')
#         else:
#             messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")

#     return render(request, 'core/list_product_reward.html', {
#         'product_rewards': product_rewards,
#         'search_query': search_query,
#         'user_points': user_points.points if user_points else 0,  # ✅ Points mis à jour
#         'spent_points': user_points.spent_points if user_points else 0,  # ✅ Points dépensés affichés
#         'paginator': paginator,
#     })

# def list_product_rewards(request):
#     product_rewards = ProductPoints.objects.all()

#     # Recherche
#     search_query = request.GET.get('search', '').strip()
#     if search_query:
#         product_rewards = product_rewards.filter(name__icontains=search_query) | product_rewards.filter(description__icontains=search_query)

#     # Pagination
#     paginator = Paginator(product_rewards, 6)
#     page = request.GET.get('page')
#     try:
#         product_rewards = paginator.page(page)
#     except PageNotAnInteger:
#         product_rewards = paginator.page(1)
#     except EmptyPage:
#         product_rewards = paginator.page(paginator.num_pages)

#     # Récupérer les points de l'utilisateur
#     try:
#         user_points = UserPoints.objects.get(user=request.user)
#     except UserPoints.DoesNotExist:
#         user_points = None

#     if request.method == 'POST':
#         product_id = request.POST.get('product_id')
#         product = ProductPoints.objects.get(id=product_id)

#         if user_points and user_points.spend_points(product.points_required):
#             # Enregistrer l'achat
#             Purchase.objects.create(
#                 user=request.user,
#                 product=product,
#                 points_used=product.points_required
#             )
#             messages.success(request, f"Achat de {product.name} effectué avec succès !")
#             return redirect('list_product_rewards')
#         else:
#             messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")

#     return render(request, 'core/list_product_reward.html', {
#         'product_rewards': product_rewards,
#         'search_query': search_query,
#         'user_points': user_points.points if user_points else 0,  # ✅ Points mis à jour
#         'paginator': paginator,
#     })

from django.shortcuts import render, get_object_or_404, redirect
from .models import ProductPoints, UserPoints, Purchase
from django.contrib import messages

def buy_product(request, product_id):
    product = get_object_or_404(ProductPoints, id=product_id)

    # Utiliser filter() pour récupérer l'enregistrement des points de l'utilisateur, ou None si non trouvé
    user_points = UserPoints.objects.filter(user=request.user).first()

    if not user_points:
        # Si l'utilisateur n'a pas de points, créer un enregistrement avec 0 points
        user_points = UserPoints.objects.create(user=request.user, points=0)

    # Vérifier si l'utilisateur a assez de points
    if user_points.points >= product.points_required:
        if request.method == 'POST':
            # Déduire les points de l'utilisateur
            user_points.points -= product.points_required
            user_points.save()

            # Enregistrer l'achat dans le modèle Purchase
            purchase = Purchase.objects.create(
                user=request.user,
                product=product,
                points_used=product.points_required
            )

            # Ajouter un message de succès
            messages.success(request, f"Achat de {product.name} effectué avec succès !")
            return redirect('succees')  # Rediriger vers la page de confirmation

        # Afficher la page de confirmation d'achat
        return render(request, 'list_product_reward.html', {'product': product, 'user_points': user_points})

    else:
        messages.error(request, "Vous n'avez pas assez de points pour effectuer cet achat.")
        return redirect('insufficient_points')  # Rediriger vers la page d'erreur si pas assez de points

# core/views.py
from django.shortcuts import render

def insufficient_points(request):
    # Afficher la page d'erreur lorsque l'utilisateur n'a pas assez de points
    return render(request, 'core/insufficient_points.html')

# core/views.py
from django.shortcuts import render

def succees(request):
    return render(request, 'core/succees.html')

# @login_required
# def orders_by_date_detail_detail(request, slug, order_date, client_username):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)
#     user = get_object_or_404(CustomUser, username=client_username)

#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

#     # Récupérer les commandes du client pour ce store et cette date
#     orders = Order.objects.filter(
#         user__username=client_username,
#         created_at__date=order_date,
#         items__product__store=store,
#         activated=True
#     ).distinct()

#     # Créer une liste de détails de commande, avec les noms des articles
#     order_data = []
#     for order in orders:
#         # Récupérer tous les items associés à cette commande
#         order_items = order.items.filter(product__store=store)

#         # Calculer le nombre d'articles pour cette commande dans ce store
#         total_items_in_order = order_items.count()

#         # Calculer le montant total des articles pour cette commande
#         total_order_amount = order_items.aggregate(total_amount=Sum('price_at_time_of_order'))['total_amount'] or 0

#         # Extraire les noms des articles pour cette commande
#         item_names = ', '.join([item.product.name for item in order_items])

#         # Ajouter les informations nécessaires
#         order_data.append({
#             'order_id': order.id,
#             'total_items_in_order': total_items_in_order,
#             'total_order_amount': total_order_amount,
#             'item_names': item_names,  # Liste des noms des articles
#         })

#     # Passer les données à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'order_data': order_data,
#         'order_date': order_date,
#         'client_username': client_username,
#         'user_phone': user.phone,
#     }

#     return render(request, 'core/orders_by_date_detail_detail.html', context)

# @login_required
# def orders_by_date_detail_detail(request, slug, order_date, client_username):
#     # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
#     store = get_object_or_404(Store, slug=slug, owner=request.user)
#     user = get_object_or_404(CustomUser, username=client_username)

#     # Convertir la date reçue dans l'URL en un objet date
#     order_date = datetime.strptime(order_date, "%Y-%m-%d").date()

#     # Récupérer les commandes du client pour ce store et cette date
#     orders = Order.objects.filter(
#         user__username=client_username,
#         created_at__date=order_date,
#         items__product__store=store,
#         activated=True
#     ).distinct()

#     # Créer une liste de détails de commande, avec les noms des articles
#     order_data = []
#     for order in orders:
#         # Récupérer tous les items associés à cette commande
#         order_items = order.items.filter(product__store=store)

#         # Calculer le nombre d'articles pour cette commande dans ce store
#         total_items_in_order = order_items.count()

#         # Calculer le montant total des articles pour cette commande
#         total_order_amount = order_items.aggregate(total_amount=Sum('price_at_time_of_order'))['total_amount'] or 0

#         # Extraire les noms des articles pour cette commande
#         item_names = ', '.join([item.product.name for item in order_items])

#         # Ajouter les informations nécessaires
#         order_data.append({
#             'order_id': order.id,
#             'total_items_in_order': total_items_in_order,
#             'total_order_amount': total_order_amount,
#             'item_names': item_names,  # Liste des noms des articles
#         })

#     # Passer les données à l'HTML pour l'affichage
#     context = {
#         'store': store,
#         'order_data': order_data,
#         'order_date': order_date,
#         'client_username': client_username,
#         'user_phone': user.phone,
#     }

#     return render(request, 'core/orders_by_date_detail_detail.html', context)





# @login_required
# def order_detail_store(request, client_username):
#     # Récupérer l'utilisateur en fonction du nom d'utilisateur
#     User = get_user_model()
#     user = get_object_or_404(User, username=client_username)

#     # Récupérer toutes les commandes de ce client
#     orders = Order.objects.filter(user=user)

#     # Pour chaque commande, récupérer les articles associés
#     order_data = []
#     for order in orders:
#         order_items = order.items.all()  # Récupère tous les articles de la commande
#         total_items_in_order = order_items.count()  # Nombre d'articles
#         total_order_amount = order.total_amount  # Montant total de la commande

#         # Ajouter les données de cette commande à la liste
#         order_data.append({
#             'order_id': order.id,
#             'status': order.status,
#             'created_at': order.created_at,
#             'total_amount': total_order_amount,
#             'total_items': total_items_in_order,
#         })

#     # Passer les données à ton template
#     context = {
#         'user': user,
#         'order_data': order_data,
#     }

#     return render(request, 'core/order_detail_store.html', context)

# def order_detail_store(request, order_id):
#     # Récupérer la commande par ID
#     order = get_object_or_404(Order, id=order_id)

#     # Récupérer les items associés à cette commande
#     order_items = order.items.all()  # Utilise 'items' grâce à 'related_name' dans la relation ForeignKey

#     # Passer les informations à ton template
#     context = {
#         'order': order,
#         'order_items': order_items,
#     }

#     # Afficher les détails de la commande dans le template
#     return render(request, 'core/order_detail_store.html', context)



# @login_required
# def order_detail_store(request, order_id):
#     # Récupérer la commande spécifique par ID
#     order = get_object_or_404(Order, id=order_id)

#     # Passer l'objet `order` au template
#     context = {
#         'order': order,
#     }

#     # Utilisation du nouveau template 'order_detail_store.html'
#     return render(request, 'core/order_detail_store.html', context)

from django.http import HttpResponse
from docx import Document
from django.shortcuts import get_object_or_404
from .models import Order, OrderItem  # Assure-toi que tu importes correctement tes modèles
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from .models import Order, OrderItem, CustomUser
from django.contrib.auth.decorators import login_required
from django.db.models import Sum
from django.http import HttpResponse
from docx import Document
from django.shortcuts import get_object_or_404
from .models import Order, OrderItem
from django.contrib.auth.decorators import login_required

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from docx import Document
from .models import Store, CustomUser, Order, OrderItem

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from docx import Document
from .models import Store, CustomUser, Order, OrderItem

@login_required
def generer_word(request, slug, order_date, client_username):
    # Récupérer le store en fonction du slug et s'assurer qu'il appartient à l'utilisateur
    store = get_object_or_404(Store, slug=slug, owner=request.user)
    user = get_object_or_404(CustomUser, username=client_username)

    # Récupérer l'ID de la commande depuis la requête
    order_id = request.GET.get("order_id")
    order = get_object_or_404(Order, id=order_id)

    # Vérifier si la commande appartient bien au store
    order_items = OrderItem.objects.filter(order=order, product__store=store)

    # Si aucun article n'appartient au store spécifié, renvoyer une erreur
    if not order_items.exists():
        return HttpResponse("Aucun article trouvé pour ce store dans cette commande.", status=404)

    # Calculer le montant total de la commande pour ce store
    total_order_amount_for_store = sum(item.get_total_price() for item in order_items)

    # Créer un document Word
    doc = Document()

    # Ajouter un titre
    doc.add_heading(f'Détail de la Commande #{order.id} - Store: {store.name}', 0)

    # Ajouter les informations de la commande
    doc.add_heading('Informations de la Commande', level=1)
    doc.add_paragraph(f"Client : {user.username}")
    doc.add_paragraph(f"Téléphone : {user.phone}")
    doc.add_paragraph(f"Date de commande : {order.created_at}")
    doc.add_paragraph(f"Total pour ce Store : {total_order_amount_for_store} CDF")

    # Ajouter les articles de la commande dans un tableau
    doc.add_heading('Articles de la commande', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes du tableau
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Produit'
    hdr_cells[1].text = 'Quantité'
    hdr_cells[2].text = 'Prix Unitaire'
    hdr_cells[3].text = 'Prix Total'

    # Ajouter chaque article du store spécifique dans le tableau
    for item in order_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.product.name
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = f"{item.price_at_time_of_order} CDF"
        row_cells[3].text = f"{item.get_total_price()} CDF"

    # Créer la réponse HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="commande_{order.id}_store_{store.slug}.docx"'

    # Sauvegarder le document dans la réponse
    doc.save(response)

    return response



# @login_required
# def generer_word(request, slug, order_date, client_username):
#     # Récupérer la commande spécifique par ID
#     order_id = request.GET.get('order_id')
#     order = get_object_or_404(Order, id=order_id)
#     order_items = OrderItem.objects.filter(order=order)  # Liste des articles dans la commande

#     # Créer un document Word
#     doc = Document()

#     # Ajouter un titre
#     doc.add_heading(f'Détail de la Commande #{order.id}', 0)

#     # Ajouter les informations de la commande
#     doc.add_heading('Informations de la Commande', level=1)
#     doc.add_paragraph(f"Client : {client_username}")
#     doc.add_paragraph(f"Numéro de téléphone : {order.user.phone}")
#     doc.add_paragraph(f"Statut : {order.status}")
#     doc.add_paragraph(f"Date de commande : {order.created_at}")
#     doc.add_paragraph(f"Montant Total : {order.get_total()} CDF")

#     # Ajouter les articles de la commande dans un tableau
#     doc.add_heading('Articles de la commande', level=1)
#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'

#     # En-têtes du tableau
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Produit'
#     hdr_cells[1].text = 'Quantité'
#     hdr_cells[2].text = 'Prix Unitaire'
#     hdr_cells[3].text = 'Prix Total'

#     # Ajouter chaque article dans le tableau
#     for item in order_items:
#         row_cells = table.add_row().cells
#         row_cells[0].text = item.product.name
#         row_cells[1].text = str(item.quantity)
#         row_cells[2].text = f"{item.price_at_time_of_order} CDF"
#         row_cells[3].text = f"{item.get_total_price()} CDF"

#     # Créer la réponse HTTP pour le téléchargement
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = f'attachment; filename="commande_{order.id}.docx"'

#     # Sauvegarder le document dans la réponse
#     doc.save(response)

#     return response




# @login_required
# def order_detail(request, order_id):
#     # Récupérer la commande spécifique par ID
#     order = get_object_or_404(Order, id=order_id)

#     # Passer l'objet `order` au template
#     context = {
#         'order': order,
#     }

#     return render(request, 'core/order_detail.html', context)


from django.shortcuts import render, redirect
from .models import Store, Testimonial,Testimonialproduct
from .forms import TestimonialForm,TestimonialproductForm

def add_testimonial(request, slug):
    store = Store.objects.get(slug=slug)

    if request.method == 'POST':
        form = TestimonialForm(request.POST)
        if form.is_valid():
            # Créer le témoignage
            testimonial = form.save(commit=False)
            testimonial.user = request.user  # Assigner l'utilisateur connecté
            testimonial.store = store  # Assigner le magasin
            testimonial.save()
            messages.success(request, 'votre témoignage a été ajoutée avec succès.')
            return redirect('store_detail', slug=store.slug)
    else:
        form = TestimonialForm()

    return render(request, 'core/add_testimonial.html', {'form': form, 'store': store})

def add_testimonialproduct(request, id):
    product = Product.objects.get(id=id)

    if request.method == 'POST':
        form = TestimonialproductForm(request.POST)
        if form.is_valid():
            # Créer le témoignage
            testimonial = form.save(commit=False)
            testimonial.user = request.user  # Assigner l'utilisateur connecté
            testimonial.product = product  # Assigner le magasin
            testimonial.save()
            messages.success(request, 'votre témoignage a été ajoutée avec succès.')
            return redirect('product_detail', id=product.id)
    else:
        form = TestimonialproductForm()

    return render(request, 'core/add_testimonialproduct.html', {'form': form, 'product': product})




from .forms import CategoryForm
from .models import Store
from django.contrib.auth.decorators import login_required

@login_required
def create_category(request, slug):
    # Récupérer le store correspondant au slug
    store = get_object_or_404(Store, slug=slug)

    if request.method == 'POST':
        form = CategoryForm(request.POST)
        if form.is_valid():
            # Ajouter la catégorie au store
            category = form.save(commit=False)
            category.store = store
            category.save()

            # Afficher un message de succès
            messages.success(request, 'La catégorie a été ajoutée avec succès.')
            return redirect('store_detail', slug=store.slug)  # Rediriger vers la page du store
    else:
        form = CategoryForm()

    return render(request, 'core/create_category.html', {'form': form, 'store': store})

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Product, Photo, Store, Category
from .forms import ProductForm, PhotoForm

def create_product(request, slug):
    store = get_object_or_404(Store, slug=slug) 
     # Récupérer le store par son slug

    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)

        if form.is_valid():
            product = form.save(commit=False)  # Ne pas encore enregistrer dans la base
            product.store = store  # Associer le store au produit
            product.save()  # Enregistrer le produit

            # Gérer l'upload des images supplémentaires (galerie)
            images = request.FILES.getlist('image_galerie')  # Récupérer les images pour la galerie
            for img in images:
                Photo.objects.create(product=product, image=img)  # Créer des objets Photo pour chaque image

            # Message de succès
            messages.success(request, f"Le produit '{product.name}' a été ajouté avec succès au store {store.name}.")
            return redirect('store_detail', slug=store.slug) # Redirection vers la page de détail du store

    else:
        form = ProductForm()

    return render(request, 'core/create_product.html', {
        'form': form,
        'store': store,
    })

from django.shortcuts import render, get_object_or_404
from .models import Product,Testimonialproduct

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render, get_object_or_404
from .models import Product, Testimonialproduct

def product_detail(request, id):
    # Récupère le produit en fonction de son ID
    product = get_object_or_404(Product, id=id)
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    # Récupère la catégorie du produit
    category = product.category

    # Récupère les autres produits de la même catégorie (en excluant le produit actuel)
    related_products = Product.objects.filter(category=category).exclude(id=product.id)
    
    # Récupère les témoignages associés au produit
    testimonials = Testimonialproduct.objects.filter(product=product)
    average_rating = testimonials.aggregate(Avg('rating'))['rating__avg']
    if not average_rating:
        average_rating = 0  # Si pas de témoignages, la note moyenne est 0

    # Pagination des témoignages
    paginator_testimonials = Paginator(testimonials, 3)  # 5 témoignages par page
    page_testimonials = request.GET.get('page')

    try:
        testimonials = paginator_testimonials.page(page_testimonials)
    except PageNotAnInteger:
        testimonials = paginator_testimonials.page(1)
    except EmptyPage:
        testimonials = paginator_testimonials.page(paginator_testimonials.num_pages)

    # Pagination des produits liés
    paginator_related = Paginator(related_products, 6)  # 6 produits par page
    page_related = request.GET.get('page')

    try:
        related_products = paginator_related.page(page_related)
    except PageNotAnInteger:
        related_products = paginator_related.page(1)
    except EmptyPage:
        related_products = paginator_related.page(paginator_related.num_pages)

    # Passe le produit et les produits associés au template
    return render(request, 'core/product_detail.html', {
        'product': product,
        'related_products': related_products,  # Les produits de la même catégorie
        'paginator_related': paginator_related,  # Pour la pagination des produits liés
        'paginator_testimonials': paginator_testimonials,  # Pour la pagination des témoignages
        'testimonials': testimonials,  # Les témoignages paginés
        'range_10': range(1, 11),
        'average_rating': average_rating, 
        'favorite_stores':favorite_stores,
    })
 

@login_required
def store_sales_history(request, store_id):
    store = Store.objects.get(id=store_id, owner=request.user)
   
    # Obtenir toutes les commandes de la boutique
    orders = Order.objects.filter(store=store).order_by('-created_at') 



@login_required
def purchase_history(request):
    # Obtenir toutes les commandes de l'utilisateur connecté
    orders = Order.objects.filter(user=request.user).order_by('-created_at')  # Tri par date décroissante
   
    return render(request, 'shop/purchase_history.html', {'orders': orders})

from django.shortcuts import render, get_object_or_404
from .models import Order

@login_required(login_url="signin")
def order_detail(request, order_id):
    # Récupérer la commande en fonction de l'ID
    order = get_object_or_404(Order, id=order_id, user=request.user)
    
    # Ajouter les éléments de la commande (OrderItem)
    order_items = order.items.all()
    
    # Contexte à passer au template
    context = {
        'order': order,
        'order_items': order_items,
    }
    
    return render(request, 'core/order_detail.html', context)


from django.shortcuts import render, redirect
from .forms import CommandeLivraisonForm
from django.contrib import messages
@login_required(login_url="signin")
def creer_commande(request):
    if request.method == 'POST':
        form = CommandeLivraisonForm(request.POST)
        if form.is_valid():
            # Sauvegarder la commande et l'associer à l'utilisateur
            commande = form.save(commit=False)
            commande.user = request.user  # Lier l'utilisateur connecté à la commande
            commande.save()
            messages.success(request, "Votre commande de livraison a été enregistrée avec succès.")
            return redirect('profile')  # Redirige vers une page de succès
        else:
            messages.error(request, "Il y a des erreurs dans votre formulaire.")
    else:
        form = CommandeLivraisonForm()

    return render(request, 'core/creer_commande.html', {'form': form})

def livraison_detail(request, commande_id):
    # Récupérer la commande en utilisant l'ID de la commande
    commande = get_object_or_404(CommandeLivraison, id=commande_id)
    
    context = {
        'commande': commande
    }
    
    return render(request, 'core/livraison_detail.html', context)


from .forms import ContactProductForm


def contact_product(request, id):
    product = get_object_or_404(Product, id=id)  # Récupérer le produit par son id

    if request.method == 'POST':
        form = ContactProductForm(request.POST)
        if form.is_valid():
            # Sauvegarder le formulaire ou envoyer un email selon votre logique
            contact_message = form.save(commit=False)
            contact_message.product = product  # Associer le message au produit
            contact_message.save()

            # Afficher un message de succès
            messages.success(request, "Votre demande de contact a été envoyée avec succès.")

            # Rediriger vers la page de détail du produit après soumission réussie
            return redirect('product_detail', id=product.id)

    else:
        form = ContactProductForm()

    return render(request, 'core/contact_product.html', {'form': form, 'product': product})

from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from .models import Store,Advertisement,AdInteraction
from .forms import ContactStoreForm

def contact_store(request, slug):
    # Récupérer le magasin en fonction du slug
    store = get_object_or_404(Store, slug=slug)

    if request.method == 'POST':
        form = ContactStoreForm(request.POST)
        if form.is_valid():
            # Sauvegarder le formulaire ou envoyer un email selon votre logique
            contact_message = form.save(commit=False)
            contact_message.store = store  # Associer le message au magasin
            contact_message.save()

            # Afficher un message de succès
            messages.success(request, "Votre demande de contact a été envoyée avec succès.")

            # Rediriger vers la page de détail du magasin après soumission réussie
            return redirect('store_detail', slug=store.slug)

    else:
        form = ContactStoreForm()

    return render(request, 'core/contact_store.html', {'form': form, 'store': store})



from django.shortcuts import get_object_or_404
from django.http import JsonResponse

def interact_with_ad(request, ad_id, interaction_type):
    if request.method == "POST":
        ad = get_object_or_404(Advertisement, id=ad_id)
        user = request.user

        # Vérifier si l'utilisateur a déjà fait cette interaction
        if AdInteraction.objects.filter(user=user, ad=ad, interaction_type=interaction_type).exists():
            return JsonResponse({"message": "Interaction déjà effectuée"}, status=400)

        # Ajouter l'interaction
        AdInteraction.objects.create(user=user, ad=ad, interaction_type=interaction_type)

        # Mettre à jour les compteurs
        if interaction_type == "like":
            ad.likes_count += 1
            points = 1
        elif interaction_type == "comment":
            ad.comments_count += 1
            points = 2
        elif interaction_type == "share":
            ad.shares_count += 1
            points = 3
        else:
            return JsonResponse({"message": "Interaction non valide"}, status=400)

        ad.save()

        # Ajouter les points à l'utilisateur
        user_points, created = UserPoints.objects.get_or_create(user=user)
        user_points.points += points
        user_points.save()

        return JsonResponse({"message": "Interaction enregistrée avec succès"}, status=200)

    return JsonResponse({"message": "Méthode non autorisée"}, status=405)


def reward_ad_interaction(user, ad, interaction_type):
    """Ajoute des points en fonction de l'interaction avec la publicité."""
    points_mapping = {'like': 1, 'comment': 2, 'share': 3}
    points = points_mapping.get(interaction_type, 0)

    if points > 0:
        user_points, created = UserPoints.objects.get_or_create(user=user)

        # Vérifier si l'utilisateur a déjà effectué cette interaction sur cette pub
        existing_interaction = AdInteraction.objects.filter(user=user, ad=ad, interaction_type=interaction_type).exists()

        if not existing_interaction:
            user_points.add_ad_points(points)
            AdInteraction.objects.create(user=user, ad=ad, interaction_type=interaction_type, points_earned=points)
            return True
    return False


from django.shortcuts import render, get_object_or_404, redirect
from .models import Advertisement, AdInteraction, UserPoints
from .forms import AdInteractionForm
from django.shortcuts import render, redirect
from .models import Advertisement, UserPoints, AdInteraction
from .forms import AdInteractionForm
from django.utils.timezone import now

from django.shortcuts import render, redirect
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db import transaction
from django.db.models import Q
from .models import Advertisement, PopUpAdvertisement, AdInteraction, UserPoints
from .forms import AdInteractionForm
from django.shortcuts import render, redirect
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from .models import Advertisement, AdInteraction, PopUpAdvertisement, UserPoints
from .forms import AdInteractionForm
from django.db.models import Q
import random
from django.db import transaction
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.core.cache import cache
from django.shortcuts import render, redirect, get_object_or_404
from django.db import transaction
from django.contrib import messages

from .models import Advertisement, PopUpAdvertisement, AdInteraction, UserLocation, UserPoints
from .forms import AdInteractionForm

def advertisement_list(request):
    ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    range_10 = range(1, 11)
    
    ads = Advertisement.objects.filter(is_active=True).order_by('-created_at')

    if request.user.is_authenticated:
        user = request.user
        user_points = UserPoints.objects.filter(user=user).first()
        if not user_points:
            user_points = UserPoints.objects.create(user=user, points=0)

        user_sex = getattr(user, 'sex', None)
        user_commune = (getattr(user, 'commune', '') or '').lower()
        user_city = (getattr(user, 'city', '') or '').lower()
        user_address = (getattr(user, 'address', '') or '').lower().split()
        user_keywords = (getattr(user, 'interests', '') or '').lower().split(',')
        user_address_str = ' '.join(user_address)  # chaîne d'adresse en minuscule
        user_keywords = [k.strip() for k in (user.interests or '').lower().split(',')]
        user_location = UserLocation.objects.filter(user=user).first()
        user_lat = float(user_location.latitude) if user_location and user_location.latitude else None
        user_lon = float(user_location.longitude) if user_location and user_location.longitude else None

        filtered_ads = []

        for ad in ads:
            show_to_user = False

            if ad.target_all_users:
                show_to_user = True
            else:
                match_count = 0
                total_conditions = 0

                if ad.target_sex:
                    total_conditions += 1
                    if ad.target_sex == user_sex:
                        match_count += 1

                if ad.target_communes:
                    total_conditions += 1
                    ad_communes = [c.lower() for c in ad.target_communes]
                    if user_commune in ad_communes:
                        match_count += 1

                if ad.target_cities:
                    total_conditions += 1
                    ad_cities = [c.lower() for c in ad.target_cities]
                    if user_city in ad_cities:
                        match_count += 1

                # if ad.target_keywords:
                #     total_conditions += 1
                #     ad_keywords = [k.strip().lower() for k in ad.target_keywords.split(',')]
                #     if any(k in user_keywords for k in ad_keywords):
                #         match_count += 1

                # if ad.target_address_keywords:
                #     total_conditions += 1
                #     ad_address_keywords = [k.strip().lower() for k in ad.target_address_keywords.split(',')]
                #     if any(k in user_address for k in ad_address_keywords):
                #         match_count += 1
                if ad.target_address_keywords:
                     total_conditions += 1
                     ad_address_keywords = [k.strip().lower() for k in ad.target_address_keywords.split(',')]
                     if any(k in user_address_str for k in ad_address_keywords):
                       match_count += 1

                if ad.target_keywords:
                     total_conditions += 1
                     ad_keywords = [k.strip().lower() for k in ad.target_keywords.split(',')]
                     if any(k in user_keywords for k in ad_keywords):
                        match_count += 1
                if ad.target_latitude and ad.target_longitude and ad.target_radius_km and user_lat and user_lon:
                    total_conditions += 1
                    distance = ((user_lat - float(ad.target_latitude))**2 + (user_lon - float(ad.target_longitude))**2) ** 0.5 * 111
                    if distance <= float(ad.target_radius_km):
                        match_count += 1

                if total_conditions > 0 and match_count == total_conditions:
                    show_to_user = True
           
            if show_to_user:
                if ad.max_target_users:
                    with transaction.atomic():
                        ad = Advertisement.objects.select_for_update().get(id=ad.id)
                        if ad.targeted_users.count() >= ad.max_target_users:
                            show_to_user = False
                        elif not ad.targeted_users.filter(id=user.id).exists():
                            ad.targeted_users.add(user)
                elif not ad.targeted_users.filter(id=user.id).exists():
                    ad.targeted_users.add(user)

            ad.user_has_liked = AdInteraction.objects.filter(
                user=user,
                ad=ad,
                interaction_type='like'
            ).exists()

            if show_to_user:
                filtered_ads.append(ad)

        ads = filtered_ads

    else:
        user_points = None
        for ad in ads:
            ad.user_has_liked = False
        ads = [ad for ad in ads if ad.target_all_users]

    if request.method == 'POST':
        form = AdInteractionForm(request.POST)
        if form.is_valid():
            ad = form.cleaned_data['ad']
            interaction_type = form.cleaned_data['interaction_type']

            if request.user.is_authenticated and not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
                AdInteraction.objects.create(
                    user=request.user,
                    ad=ad,
                    interaction_type=interaction_type
                )

                points_to_add = {
                    'like': 1,
                    'comment': 2,
                    'share': 3
                }.get(interaction_type, 0)

                user_points.add_ad_points(points_to_add)

                return redirect('advertisement_list')
    else:
        form = AdInteractionForm()
       
    user_shares = []
    ad_absolute_urls = {}

    if request.user.is_authenticated:
        # Récupère toutes les pubs déjà partagées par l'utilisateur
        shared_ads = Share.objects.filter(user=request.user, ad__in=ads).values_list('ad_id', flat=True)
        user_shares = list(shared_ads)
   
    # Génère les URLs absolues pour chaque publicité
    ad_absolute_urls = {ad.id: request.build_absolute_uri(ad.get_absolute_url()) for ad in ads}


    paginator = Paginator(ads, 6)
    page = request.GET.get('page')
    try:
        ads = paginator.page(page)
    except PageNotAnInteger:
        ads = paginator.page(1)
    except EmptyPage:
        ads = paginator.page(paginator.num_pages)

    no_ads_message = ""
    if not ads:
        if request.user.is_authenticated:
            no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
        else:
            no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

    return render(request, 'core/advertisement_list.html', {
        'ads': ads,
        'form': form,
        'user_points': user_points,
        'ad_popup': ad_popup,
        'no_ads_message': no_ads_message,
        'user_shares': user_shares,
        'ad_absolute_urls': ad_absolute_urls,
        'favorite_stores':favorite_stores,
        'range_10': range_10, 
    })

# def advertisement_list(request):
#     ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
#     ads = Advertisement.objects.all().order_by('-created_at')

#     if request.user.is_authenticated:
#         user = request.user
#         user_points = UserPoints.objects.filter(user=user).first()
#         if not user_points:
#             user_points = UserPoints.objects.create(user=user, points=0)

#         user_sex = getattr(user, 'sex', None)
#         user_commune = (getattr(user, 'commune', '') or '').lower()
#         user_city = (getattr(user, 'city', '') or '').lower()
#         user_address = (getattr(user, 'address', '') or '').lower().split()
#         user_keywords = (getattr(user, 'interests', '') or '').lower().split(',')

#         user_location = UserLocation.objects.filter(user=user).first()
#         user_lat = float(user_location.latitude) if user_location and user_location.latitude else None
#         user_lon = float(user_location.longitude) if user_location and user_location.longitude else None

#         filtered_ads = []

#         for ad in ads:
#             show_to_user = False

#             if ad.target_all_users:
#                 show_to_user = True
#             else:
#                 match_count = 0
#                 total_conditions = 0

#                 # Sexe ciblé
#                 if ad.target_sex:
#                     total_conditions += 1
#                     if ad.target_sex == user_sex:
#                         match_count += 1

#                 # Commune ciblée
#                 if ad.target_communes:
#                     total_conditions += 1
#                     ad_communes = [c.lower() for c in ad.target_communes]
#                     if user_commune in ad_communes:
#                         match_count += 1

#                 # Ville ciblée
#                 if ad.target_cities:
#                     total_conditions += 1
#                     ad_cities = [c.lower() for c in ad.target_cities]
#                     if user_city in ad_cities:
#                         match_count += 1

#                 # Mots-clés ciblés
#                 if ad.target_keywords:
#                     total_conditions += 1
#                     ad_keywords = [k.strip().lower() for k in ad.target_keywords.split(',')]
#                     if any(k in user_keywords for k in ad_keywords):
#                         match_count += 1

#                 # Adresse ciblée (sous forme de mots-clés dans une phrase)
#                 if ad.target_address_keywords:
#                     total_conditions += 1
#                     ad_address_keywords = [k.strip().lower() for k in ad.target_address_keywords.split(',')]
#                     if any(k in user_address for k in ad_address_keywords):
#                         match_count += 1

#                 # Ciblage géographique par distance
#                 if ad.target_latitude and ad.target_longitude and ad.target_radius_km and user_lat and user_lon:
#                     total_conditions += 1
#                     distance = ((user_lat - float(ad.target_latitude))**2 + (user_lon - float(ad.target_longitude))**2) ** 0.5 * 111  # approx conversion deg -> km
#                     if distance <= float(ad.target_radius_km):
#                         match_count += 1

#                 if total_conditions > 0 and match_count == total_conditions:
#                     show_to_user = True

#             if show_to_user:
#                 if ad.max_target_users:
#                     with transaction.atomic():
#                         ad = Advertisement.objects.select_for_update().get(id=ad.id)
#                         if ad.targeted_users.count() >= ad.max_target_users:
#                             show_to_user = False
#                         elif not ad.targeted_users.filter(id=user.id).exists():
#                             ad.targeted_users.add(user)
#                 elif not ad.targeted_users.filter(id=user.id).exists():
#                     ad.targeted_users.add(user)

#             ad.user_has_liked = AdInteraction.objects.filter(
#                 user=user,
#                 ad=ad,
#                 interaction_type='like'
#             ).exists()

#             if show_to_user:
#                 filtered_ads.append(ad)

#         ads = filtered_ads

#     else:
#         user_points = None
#         for ad in ads:
#             ad.user_has_liked = False
#         ads = [ad for ad in ads if ad.target_all_users]

#     if request.method == 'POST':
#         form = AdInteractionForm(request.POST)
#         if form.is_valid():
#             ad = form.cleaned_data['ad']
#             interaction_type = form.cleaned_data['interaction_type']

#             if request.user.is_authenticated and not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
#                 AdInteraction.objects.create(
#                     user=request.user,
#                     ad=ad,
#                     interaction_type=interaction_type
#                 )

#                 points_to_add = {
#                     'like': 1,
#                     'comment': 2,
#                     'share': 3
#                 }.get(interaction_type, 0)

#                 user_points.add_ad_points(points_to_add)

#                 return redirect('advertisement_list')
#     else:
#         form = AdInteractionForm()

#     paginator = Paginator(ads, 6)
#     page = request.GET.get('page')
#     try:
#         ads = paginator.page(page)
#     except PageNotAnInteger:
#         ads = paginator.page(1)
#     except EmptyPage:
#         ads = paginator.page(paginator.num_pages)

#     no_ads_message = ""
#     if not ads:
#         if request.user.is_authenticated:
#             no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
#         else:
#             no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

#     return render(request, 'core/advertisement_list.html', {
#         'ads': ads,
#         'form': form,
#         'user_points': user_points,
#         'ad_popup': ad_popup,
#         'no_ads_message': no_ads_message
#     })

# def advertisement_list(request):
#     ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
#     ads = Advertisement.objects.all().order_by('-created_at')

#     if request.user.is_authenticated:
#         user = request.user
#         user_points = UserPoints.objects.filter(user=user).first()
#         if not user_points:
#             user_points = UserPoints.objects.create(user=user, points=0)

#         user_sex = getattr(user, 'sex', None)
#         user_commune = (getattr(user, 'commune', '') or '').lower()
#         user_city = (getattr(user, 'city', '') or '').lower()
#         user_keywords = (getattr(user, 'interests', '') or '').lower().split(',')

#         filtered_ads = []

#         for ad in ads:
#             show_to_user = False

#             if ad.target_all_users:
#                 show_to_user = True
#             else:
#                 match_count = 0
#                 total_conditions = 0

#                 # Sexe ciblé
#                 if ad.target_sex:
#                     total_conditions += 1
#                     if ad.target_sex == user_sex:
#                         match_count += 1

#                 # Commune ciblée
#                 if ad.target_communes:
#                     total_conditions += 1
#                     ad_communes = [c.lower() for c in ad.target_communes]
#                     if user_commune in ad_communes:
#                         match_count += 1

#                 # Ville ciblée
#                 if ad.target_cities:
#                     total_conditions += 1
#                     ad_cities = [c.lower() for c in ad.target_cities]
#                     if user_city in ad_cities:
#                         match_count += 1

#                 # Mots-clés ciblés
#                 if ad.target_keywords:
#                     total_conditions += 1
#                     ad_keywords = [k.strip().lower() for k in ad.target_keywords.split(',')]
#                     if any(k in user_keywords for k in ad_keywords):
#                         match_count += 1

#                 # Si toutes les conditions définies sont remplies, on autorise
#                 if total_conditions > 0 and match_count == total_conditions:
#                     show_to_user = True

#             if show_to_user:
#                 if ad.max_target_users:
#                     with transaction.atomic():
#                         ad = Advertisement.objects.select_for_update().get(id=ad.id)
#                         if ad.targeted_users.count() >= ad.max_target_users:
#                             show_to_user = False
#                         elif not ad.targeted_users.filter(id=user.id).exists():
#                             ad.targeted_users.add(user)
#                 elif not ad.targeted_users.filter(id=user.id).exists():
#                     ad.targeted_users.add(user)

#             ad.user_has_liked = AdInteraction.objects.filter(
#                 user=user,
#                 ad=ad,
#                 interaction_type='like'
#             ).exists()

#             if show_to_user:
#                 filtered_ads.append(ad)

#         ads = filtered_ads

#     else:
#         user_points = None
#         for ad in ads:
#             ad.user_has_liked = False
#         ads = [ad for ad in ads if ad.target_all_users]

#     if request.method == 'POST':
#         form = AdInteractionForm(request.POST)
#         if form.is_valid():
#             ad = form.cleaned_data['ad']
#             interaction_type = form.cleaned_data['interaction_type']

#             if request.user.is_authenticated and not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
#                 AdInteraction.objects.create(
#                     user=request.user,
#                     ad=ad,
#                     interaction_type=interaction_type
#                 )

#                 points_to_add = {
#                     'like': 1,
#                     'comment': 2,
#                     'share': 5
#                 }.get(interaction_type, 0)

#                 user_points.add_ad_points(points_to_add)

#                 return redirect('advertisement_list')
#     else:
#         form = AdInteractionForm()

#     paginator = Paginator(ads, 6)
#     page = request.GET.get('page')
#     try:
#         ads = paginator.page(page)
#     except PageNotAnInteger:
#         ads = paginator.page(1)
#     except EmptyPage:
#         ads = paginator.page(paginator.num_pages)

#     no_ads_message = ""
#     if not ads:
#         if request.user.is_authenticated:
#             no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
#         else:
#             no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

#     return render(request, 'core/advertisement_list.html', {
#         'ads': ads,
#         'form': form,
#         'user_points': user_points,
#         'ad_popup': ad_popup,
#         'no_ads_message': no_ads_message
#     })

# def advertisement_list(request):
#     ad_popup = PopUpAdvertisement.objects.filter(is_active=True).first()
#     ads = Advertisement.objects.all().order_by('-created_at')

#     if request.user.is_authenticated:
#         user = request.user
#         user_points = UserPoints.objects.filter(user=user).first()
#         if not user_points:
#             user_points = UserPoints.objects.create(user=user, points=0)

#         user_sex = getattr(user, 'sex', None)
#         user_commune = (getattr(user, 'commune', '') or '').lower()
#         user_city = (getattr(user, 'city', '') or '').lower()
#         user_keywords = (getattr(user, 'interests', '') or '').lower().split(',')

#         filtered_ads = []

#         for ad in ads:
#             show_to_user = False

#             if ad.target_all_users:
#                 show_to_user = True
#             else:
#                 match_count = 0
#                 total_conditions = 0

#                 # Sexe ciblé
#                 if ad.target_sex:
#                     total_conditions += 1
#                     if ad.target_sex == user_sex:
#                         match_count += 1

#                 # Commune ciblée
#                 if ad.target_communes:
#                     total_conditions += 1
#                     ad_communes = [c.lower() for c in ad.target_communes]
#                     if user_commune in ad_communes:
#                         match_count += 1

#                 # Ville ciblée
#                 if ad.target_cities:
#                     total_conditions += 1
#                     ad_cities = [c.lower() for c in ad.target_cities]
#                     if user_city in ad_cities:
#                         match_count += 1

#                 # Mots-clés ciblés
#                 if ad.target_keywords:
#                     total_conditions += 1
#                     ad_keywords = [k.strip().lower() for k in ad.target_keywords.split(',')]
#                     if any(k in user_keywords for k in ad_keywords):
#                         match_count += 1

#                 # Si au moins une condition est remplie (ou toutes si définies), on autorise
#                 if total_conditions > 0 and match_count == total_conditions:
#                     show_to_user = True

#             if show_to_user:
#                 # Respect du nombre max d'utilisateurs ciblés
#                 if ad.max_target_users and ad.targeted_users.count() >= ad.max_target_users:
#                     show_to_user = False
#                 elif not ad.targeted_users.filter(id=user.id).exists():
#                     ad.targeted_users.add(user)

#             ad.user_has_liked = AdInteraction.objects.filter(
#                 user=user,
#                 ad=ad,
#                 interaction_type='like'
#             ).exists()

#             if show_to_user:
#                 filtered_ads.append(ad)

#         ads = filtered_ads

#     else:
#         user_points = None
#         for ad in ads:
#             ad.user_has_liked = False
#         ads = [ad for ad in ads if ad.target_all_users]

#     if request.method == 'POST':
#         form = AdInteractionForm(request.POST)
#         if form.is_valid():
#             ad = form.cleaned_data['ad']
#             interaction_type = form.cleaned_data['interaction_type']

#             if request.user.is_authenticated and not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
#                 AdInteraction.objects.create(
#                     user=request.user,
#                     ad=ad,
#                     interaction_type=interaction_type
#                 )

#                 points_to_add = {
#                     'like': 1,
#                     'comment': 2,
#                     'share': 5
#                 }.get(interaction_type, 0)

#                 user_points.add_ad_points(points_to_add)

#                 return redirect('advertisement_list')
#     else:
#         form = AdInteractionForm()

#     paginator = Paginator(ads, 6)
#     page = request.GET.get('page')
#     try:
#         ads = paginator.page(page)
#     except PageNotAnInteger:
#         ads = paginator.page(1)
#     except EmptyPage:
#         ads = paginator.page(paginator.num_pages)

#     no_ads_message = ""
#     if not ads:
#         if request.user.is_authenticated:
#             no_ads_message = "Aucune publicité ne correspond à votre profil pour le moment."
#         else:
#             no_ads_message = "Aucune publicité disponible actuellement pour tous les utilisateurs."

#     return render(request, 'core/advertisement_list.html', {
#         'ads': ads,
#         'form': form,
#         'user_points': user_points,
#         'ad_popup': ad_popup,
#         'no_ads_message': no_ads_message
#     })



# def advertisement_list(request):
#     # Récupère toutes les publicités
#     ads = Advertisement.objects.all()

#     # Récupérer les points de l'utilisateur pour l'affichage dans la vue (optionnel)
#     user_points = UserPoints.objects.filter(user=request.user).first()
#     if not user_points:
#         user_points = UserPoints.objects.create(user=request.user, points=0)

#     # Gérer les interactions
#     if request.method == 'POST':
#         form = AdInteractionForm(request.POST)
#         if form.is_valid():
#             ad = form.cleaned_data['ad']
#             interaction_type = form.cleaned_data['interaction_type']
            
#             # Vérifier si l'utilisateur a déjà effectué cette interaction
#             if not AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type=interaction_type).exists():
#                 # Créer l'interaction
#                 AdInteraction.objects.create(
#                     user=request.user,
#                     ad=ad,
#                     interaction_type=interaction_type
#                 )
                
#                 # Ajouter des points pour l'utilisateur
#                 points_to_add = 0
#                 if interaction_type == 'like':
#                     points_to_add = 1
#                 elif interaction_type == 'comment':
#                     points_to_add = 2
#                 elif interaction_type == 'share':
#                     points_to_add = 5

#                 # Ajouter les points à l'utilisateur
#                 user_points.add_ad_points(points_to_add)
                
#                 # Rediriger pour éviter la soumission multiple
#                 return redirect('advertisement_list')

#     else:
#         form = AdInteractionForm()

#     return render(request, 'core/advertisement_list.html', {
#         'ads': ads,
#         'form': form,
#         'user_points': user_points
#     })



from django.shortcuts import render
from .models import Advertisement

from django.shortcuts import render, get_object_or_404
from .models import Advertisement, Share

from django.shortcuts import render, get_object_or_404, redirect
from .models import Advertisement, Share, UserPoints

from django.shortcuts import get_object_or_404, redirect

def advertisement_detail(request, slug):
    ad = Advertisement.objects.filter(slug=slug).first()
    favorite_stores = Store.objects.filter(favoritestore=True).order_by('-created_at')
    range_10 = range(1, 11)
    # Si l'annonce n'existe pas, rediriger vers la liste des publicités
    if not ad:
        return redirect('advertisement_list')

    ad_absolute_url = request.build_absolute_uri(ad.get_absolute_url())

    # Vérifie si l'utilisateur est authentifié avant de vérifier le partage
    user_shared = False
    if request.user.is_authenticated:
        user_shared = Share.objects.filter(user=request.user, ad=ad).exists()

    return render(request, 'core/advertisement_detail.html', {
        'ad': ad,
        'user_shared': user_shared,
        'ad_absolute_url': ad_absolute_url,
        'favorite_stores':favorite_stores,
        'range_10': range_10, 
    })


# views.py
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Advertisement, AdInteraction, UserPoints
# views.py

from django.contrib import messages  # Importation des messages
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from .models import Advertisement, AdInteraction, UserPoints

def visit_ad_url(request, slug):
    ad = get_object_or_404(Advertisement, slug=slug)

    # Incrémente le compteur de visites
    ad.visits_count += 1
    ad.save()

    if request.user.is_authenticated:
        # Vérifie si l'utilisateur a déjà gagné le point bonus via l'autre vue
        already_got_bonus_point = AdInteraction.objects.filter(
            user=request.user,
            ad=ad,
            interaction_type='bonus_1_point'
        ).exists()

        if already_got_bonus_point:
            messages.info(request, "Vous avez déjà gagné le point bonus pour cette publicité. Vous ne pouvez pas gagner un autre point ici.")
            return redirect(ad.url)

        # Vérifie s’il a déjà eu le point ici
        already_visited = AdInteraction.objects.filter(
            user=request.user,
            ad=ad,
            interaction_type='visit'
        ).exists()

        if already_visited:
            messages.info(request, "Vous avez déjà gagné 1 point pour cette visite.")
        else:
            # Crée l’interaction et ajoute 1 point
            AdInteraction.objects.create(
                user=request.user,
                ad=ad,
                interaction_type='visit'
            )

            user_points, _ = UserPoints.objects.get_or_create(user=request.user)
            user_points.points += 1
            user_points.ad_points += 1
            user_points.save()

            messages.success(request, "Vous avez gagné 1 point pour avoir visité cette publicité !")
    else:
        messages.info(request, "Connecte-toi pour gagner des points en visitant les liens des publicités.")

    return redirect(ad.url)



from django.db.models import Q

def visit_ad_url_no_points(request, slug):
    ad = get_object_or_404(Advertisement, slug=slug)

    # Incrémenter le compteur de visites
    ad.visits_count += 1
    ad.save()

    if request.user.is_authenticated:
        # Vérifier si l'utilisateur a déjà eu les 2 points pour cette pub
        has_visited_main = AdInteraction.objects.filter(
            user=request.user,
            ad=ad,
            interaction_type='visit'
        ).exists()

        if has_visited_main:
            messages.info(request, "Vous avez déjà gagné 1 point pour cette publicité. Vous ne pouvez pas gagner ce point supplémentaire.")
            return redirect(ad.url)

        # Vérifier si quelqu’un a déjà pris le point unique
        point_already_taken = AdInteraction.objects.filter(
            ad=ad,
            interaction_type='bonus_1_point'
        ).exists()

        if point_already_taken:
           messages.info(request, "Le point bonus de cette publicité a déjà été remporté par un autre utilisateur. Continue de visiter les liens des publicités, ta régularité sera bientôt récompensée par l’administration !")
        else:
            # Créer une interaction de type "bonus_1_point"
            AdInteraction.objects.create(
                user=request.user,
                ad=ad,
                interaction_type='bonus_1_point'
            )

            # Ajouter 1 point au profil de l'utilisateur
            user_points, _ = UserPoints.objects.get_or_create(user=request.user)
            user_points.points += 1
            user_points.ad_points += 1
            user_points.save()

            messages.success(request, "Félicitations ! Vous avez gagné 1 point bonus en étant le premier à visiter ce lien.")
    else:
        messages.info(request, "Connecte-toi pour tenter de gagner le point bonus réservé au premier visiteur.")

    return redirect(ad.url)

# def visit_ad_url_no_points(request, slug):
#     ad = get_object_or_404(Advertisement, slug=slug)

#     # Incrémente juste le compteur de visites
#     ad.visits_count += 1
#     ad.save()

#     return redirect(ad.url)

# @login_required
# def visit_ad_url(request, slug):
#     ad = get_object_or_404(Advertisement, slug=slug)

#     already_visited = AdInteraction.objects.filter(
#         user=request.user,
#         ad=ad,
#         interaction_type='visit'
#     ).exists()

#     if already_visited:
#         # Si l'utilisateur a déjà visité cette publicité
#         messages.info(request, "Vous avez déjà gagné 2 points pour cette visite.")
#     else:
#         # Crée l'interaction pour la première visite
#         AdInteraction.objects.create(
#             user=request.user,
#             ad=ad,
#             interaction_type='visit'
#         )

#         # Ajoute les points
#         user_points, _ = UserPoints.objects.get_or_create(user=request.user)
#         user_points.points += 2
#         user_points.ad_points += 2
#         user_points.save()

#         # Incrémente le compteur de visites
#         ad.visits_count += 1
#         ad.save()

#         # Message de succès
#         messages.success(request, "Vous avez gagné 2 points pour avoir visité cette publicité !")

#     return redirect(ad.url)


# def advertisement_detail(request, slug):
#     ad = Advertisement.objects.get(slug=slug)
#     ad_absolute_url = request.build_absolute_uri(ad.get_absolute_url())

#     # Vérifie si l'utilisateur est authentifié avant de vérifier le partage
#     user_shared = False
#     if request.user.is_authenticated:
#         user_shared = Share.objects.filter(user=request.user, ad=ad).exists()

#     return render(request, 'core/advertisement_detail.html', {
#         'ad': ad,
#         'user_shared': user_shared,
#         'ad_absolute_url': ad_absolute_url
#     })

# def advertisement_detail(request, slug):
#     ad = Advertisement.objects.get(slug=slug)

#     # Obtenir l'URL absolue
#     ad_absolute_url = request.build_absolute_uri(ad.get_absolute_url())

#     # Vérifie si l'utilisateur a déjà partagé cette publicité
#     user_shared = Share.objects.filter(user=request.user, ad=ad).exists()

#     return render(request, 'core/advertisement_detail.html', {
#         'ad': ad,
#         'user_shared': user_shared,
#         'ad_absolute_url': ad_absolute_url  # Ajouter l'URL absolue de la publicité
#     })




from django.shortcuts import redirect
from django.shortcuts import redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from .models import Advertisement, Share, UserPoints

from django.shortcuts import redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import Advertisement, Share, UserPoints
from django.http import JsonResponse
from .models import Advertisement, Share



@login_required
def record_share(request, slug):
    ad = Advertisement.objects.get(slug=slug)

    # Vérifie si l'utilisateur a déjà partagé cette publicité
    if not Share.objects.filter(user=request.user, ad=ad).exists():
        # Crée une nouvelle entrée de partage
        Share.objects.create(user=request.user, ad=ad)
        
        # Incrémente le compteur de partages
        ad.shares_count += 1
        ad.save()

        # Ajoute 5 points à l'utilisateur
        user_points = UserPoints.objects.get(user=request.user)
        user_points.points += 3
        user_points.save()

        # Retourner une réponse JSON pour la mise à jour dynamique
        return JsonResponse({
            'status': 'success',
            'shares_count': ad.shares_count,
            'points': user_points.points
        })

    return JsonResponse({'status': 'error', 'message': 'Déjà partagé'})


from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Advertisement, Share, UserPoints

@login_required
def ad_share(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug)

    # Si la pub est déjà désactivée, empêcher toute action
    if not ad.is_active:
        messages.error(request, "Cette publicité n'est plus disponible.")
        return redirect('core:advertisement_detail', slug=ad_slug)

    # Vérifier si le nombre maximal de partages est atteint
    if ad.max_shares is not None and ad.shares_count >= ad.max_shares:
        ad.is_active = False
        ad.save()
        messages.error(request, "Le nombre maximal de partages est atteint. Cette publicité est maintenant désactivée.")
        return redirect('core:advertisement_detail', slug=ad_slug)

    # Vérifier si l'utilisateur a déjà partagé cette publicité
    if Share.objects.filter(user=request.user, ad=ad).exists():
        messages.error(request, "Vous avez déjà partagé cette publicité.")
    else:
        # Créer une nouvelle entrée de partage
        Share.objects.create(user=request.user, ad=ad)

        # Ajouter des points
        user_points = UserPoints.objects.get(user=request.user)
        user_points.points += 3
        user_points.save()

        # Incrémenter le compteur de partages
        ad.shares_count += 1

        # Vérifier à nouveau après incrémentation si on atteint la limite
        if ad.max_shares is not None and ad.shares_count >= ad.max_shares:
            ad.is_active = False  # Désactivation automatique
            messages.success(request, "Vous avez partagé cette publicité. Elle est maintenant désactivée car la limite a été atteinte.")
        else:
            messages.success(request, "Vous avez partagé cette publicité et gagné 3 points !")

        ad.save()

    return redirect('core:advertisement_detail', slug=ad_slug)

# @login_required
# def ad_share(request, ad_slug):
#     ad = get_object_or_404(Advertisement, slug=ad_slug)

#     # Vérifier si l'utilisateur a déjà partagé la publicité
#     if Share.objects.filter(user=request.user, ad=ad).exists():
#         messages.error(request, "Vous avez déjà partagé cette publicité,merci de le faire à nouveau.")
#     else:
#         # Créer une nouvelle entrée de partage
#         Share.objects.create(user=request.user, ad=ad)

#         # Ajouter les points pour le partage (5 points)
#         user_points = UserPoints.objects.get(user=request.user)
#         user_points.points += 3
#         user_points.save()

#         # Optionnel : Mettre à jour un compteur de partages dans l'objet Advertisement (si tu veux)
#         ad.shares_count += 1
#         ad.save()

#         messages.success(request, "Vous avez partagé cette publicité et gagné 3 points!")

#     return redirect('core:advertisement_detail', slug=ad_slug)



from django.shortcuts import get_object_or_404
from .models import Advertisement, AdInteraction

from django.contrib import messages
from django.contrib import messages
from django.shortcuts import get_object_or_404, redirect
from .models import Advertisement, AdInteraction, UserPoints

@login_required
def handle_like(request, ad_id):
    ad = get_object_or_404(Advertisement, id=ad_id, is_active=True)  # on ne peut liker que les pubs actives
    user = request.user

    interaction = AdInteraction.objects.filter(user=user, ad=ad).first()

    if interaction:
        if interaction.interaction_type == 'like':
            interaction.delete()
            ad.likes_count -= 1
            ad.save()

            user_points = user.userpoints
            user_points.points -= 1
            user_points.save()

            messages.success(request, "Vous avez perdu 1 point en retirant votre like (dislike).")

        else:
            interaction.interaction_type = 'like'
            interaction.save()

            ad.likes_count += 1
            ad.save()

            # ✅ Vérifie la désactivation après le like
            ad.check_deactivation_by_likes()

            user_points = user.userpoints
            user_points.points += 1
            user_points.save()

            messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

    else:
        AdInteraction.objects.create(user=user, ad=ad, interaction_type='like')
        ad.likes_count += 1
        ad.save()

        # ✅ Vérifie la désactivation après le like
        ad.check_deactivation_by_likes()

        user_points = user.userpoints
        user_points.points += 1
        user_points.save()

        messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

    return redirect('advertisement_list')

# @login_required
# def handle_like(request, ad_id):
#     ad = get_object_or_404(Advertisement, id=ad_id)
#     user = request.user

#     # Vérifier si l'utilisateur a déjà interagi avec cette publicité
#     interaction = AdInteraction.objects.filter(user=user, ad=ad).first()

#     if interaction:
#         # Si l'interaction existe déjà, gérer le "like" / "dislike"
#         if interaction.interaction_type == 'like':
#             # Si c'est un "like", on le transforme en "dislike" (on supprime le like)
#             interaction.delete()  # Supprimer le "like"
#             ad.likes_count -= 1  # Réduire le compteur de likes
#             ad.save()

#             # Réduire de 1 point l'utilisateur
#             user_points = user.userpoints
#             user_points.points -= 1
#             user_points.save()

#             # Message de notification pour la réduction des points
#             messages.success(request, "Vous avez perdu 1 point en retirant votre like (dislike).")

#         else:
#             # Si c'était un "dislike", on le transforme en "like"
#             interaction.interaction_type = 'like'
#             interaction.save()
#             ad.likes_count += 1  # Augmenter le compteur de likes
#             ad.save()

#             # Ajouter 1 point à l'utilisateur
#             user_points = user.userpoints
#             user_points.points += 1
#             user_points.save()

#             # Message de notification pour le like
#             messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

#     else:
#         # Si aucune interaction, créer un "like"
#         AdInteraction.objects.create(user=user, ad=ad, interaction_type='like')
#         ad.likes_count += 1
#         ad.save()

#         # Ajouter 1 point pour le "like"
#         user_points = user.userpoints
#         user_points.points += 1
#         user_points.save()

#         # Message de notification pour le like
#         messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

#     return redirect('advertisement_list')

# @login_required
# def handle_like(request, ad_id):
#     ad = get_object_or_404(Advertisement, id=ad_id)
#     user_interaction = AdInteraction.objects.filter(user=request.user, ad=ad, interaction_type='like').first()

#     if user_interaction:
#         user_interaction.toggle_like()  # Déclenche le toggle (like ou dislike)
#     else:
#         # Si l'utilisateur n'a pas encore aimé la pub, on ajoute un like
#         AdInteraction.objects.create(user=request.user, ad=ad, interaction_type='like')
#         ad.likes_count += 1
#         ad.save()

#         # Ajouter des points pour le like
#         user_points = UserPoints.objects.get(user=request.user)
#         user_points.points += 1  # Ajouter 1 point pour un like
#         user_points.save()

#         # Ajouter un message de notification
#         messages.success(request, "Vous avez gagné 1 point pour avoir aimé cette publicité !")

#     return redirect('advertisement_list')




# views.py
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Advertisement, Comment, AdInteraction

@login_required
def ad_comments(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug)
    comments = Comment.objects.filter(ad=ad).order_by('-created_at')

    if request.method == 'POST':
        comment_content = request.POST.get('comment_content')
        if not Comment.objects.filter(user=request.user, ad=ad).exists():
            Comment.objects.create(user=request.user, ad=ad, content=comment_content)
            ad.comments_count += 1
            ad.save()
            # Ajouter des points pour le commentaire
            user_points = UserPoints.objects.get(user=request.user)
            user_points.points += 2
            user_points.save()
        else:
            messages.error(request, "Vous avez déjà commenté cette publicité.")
        
         # Pagination : 6 commentaires par page
    # paginator = Paginator(comments, 4)  
    # page = request.GET.get('page')

    # try:
    #     comments = paginator.page(page)
    # except PageNotAnInteger:
    #     comments = paginator.page(1)
    # except EmptyPage:
    #     comments = paginator.page(paginator.num_pages)

    return render(request, 'core/ad_comments.html', {
        'ad': ad,
        'comments': comments
    })

from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.utils.http import urlencode
from .models import Advertisement, Share
from urllib.parse import urlencode
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Advertisement, Share, UserPoints
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from .models import Advertisement, AdInteraction, Share
from django.shortcuts import get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from .models import Advertisement, Share, AdInteraction, UserPoints
@login_required
def share_ad(request, ad_slug):
    ad = get_object_or_404(Advertisement, slug=ad_slug, is_active=True)
    user = request.user

    if request.method == 'POST':
        interaction, created = AdInteraction.objects.get_or_create(
            user=user, ad=ad, interaction_type='share'
        )

        if created:
            Share.objects.create(user=user, ad=ad)

            user_points = user.userpoints
            user_points.points += 3
            user_points.save()

            ad.shares_count += 1
            ad.save()

            # ✅ Désactivation si max_shares atteint
            ad.check_deactivation_by_shares()

        return JsonResponse({
            'status': 'success',
            'new_shares': ad.shares_count,
            'new_points': user_points.points
        })

    return redirect('advertisement_list')
# @login_required
# def share_ad(request, ad_slug):
#     ad = get_object_or_404(Advertisement, slug=ad_slug)
#     user = request.user

#     if request.method == 'POST':
#         # Vérifier si l'interaction de partage existe déjà
#         interaction, created = AdInteraction.objects.get_or_create(
#             user=user, ad=ad, interaction_type='share'
#         )

#         if created:  # Si l'interaction est nouvelle (partage non effectué auparavant)
#             # Ajouter un enregistrement dans le modèle Share
#             Share.objects.create(user=user, ad=ad)

#             # Ajouter des points à l'utilisateur
#             user_points = user.userpoints
#             user_points.points += 3  # Ajouter 5 points pour le partage
#             user_points.save()

#             # Incrémenter le compteur de partages de l'annonce
#             ad.shares_count += 1
#             ad.save()

#         return JsonResponse({
#             'status': 'success',
#             'new_shares': ad.shares_count,
#             'new_points': user_points.points
#         })

#     return redirect('advertisement_list')







from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from .models import AdInteraction, Advertisement, UserPoints
from django.db.models import Sum
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from core.models import UserPoints, Purchase
from django.db.models import Sum
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from core.models import UserPoints, Purchase,PointConversion,PopUpAdvertisement

@login_required(login_url="signin")
def user_dashboard(request):
    try:
        # Récupérer les points de l'utilisateur
        user_points = UserPoints.objects.get(user=request.user)
    except UserPoints.DoesNotExist:
        user_points = None

    # Récupérer les achats effectués par l'utilisateur
    user_purchases = Purchase.objects.filter(user=request.user).order_by('-purchase_date')  # Dernier achat en premier

    # Calculer le total des points dépensés
    total_spent_points = user_purchases.aggregate(total_spent=Sum('points_used'))['total_spent'] or 0
    
    # Récupérer le taux de conversion actuel
    conversion_rate = PointConversion.objects.first() 
     # Si vous avez plusieurs taux, ajustez cette logique
    if conversion_rate:
        usd_value = conversion_rate.convert_points_to_usd(user_points.points)
    else:
        usd_value = 0  # En cas d'absence de taux de conversion
    # Pagination pour les achats
    paginator = Paginator(user_purchases, 4)  # 4 achats par page
    page_number = request.GET.get('page')
    
    try:
        purchases = paginator.page(page_number)
    except PageNotAnInteger:
        # Si la page n'est pas un entier, afficher la première page
        purchases = paginator.page(1)
    except EmptyPage:
        # Si la page est vide, afficher la dernière page
        purchases = paginator.page(paginator.num_pages)

    # Calculer le total des points disponibles
    if user_points:
        total_points = user_points.points
    else:
        total_points = 0
    
     # Récupérer le taux de conversion (par exemple 1 point = 0.5 USD)
    try:
        conversion = PointConversion.objects.latest('id')  # Récupère le dernier taux de conversion ajouté
        conversion_rate = conversion.conversion_rate
    except PointConversion.DoesNotExist:
        conversion_rate = 0.5  # Valeur par défaut si aucun taux de conversion n'est défini

    # Calculer la valeur en USD des points
    total_in_usd = total_points * conversion_rate

    context = {
        'user_points': user_points,
        'total_spent_points': total_spent_points,
        'total_points': total_points,
        'purchases': purchases,  # Achats paginés
        'usd_value': usd_value,
        'total_in_usd': total_in_usd,
        
    }

    return render(request, 'core/user_dashboard.html', context)

from django.contrib.auth.decorators import user_passes_test
from django.shortcuts import render
from .utils import count_online_users

@user_passes_test(lambda u: u.is_authenticated and u.is_staff, login_url='/')
def statistiques_view(request):
    online_users = count_online_users()
    return render(request, 'core/statistiques.html', {'online_users': online_users})
# views.py



from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseRedirect
from .models import Lottery, LotteryParticipation
import random
from django.contrib.admin.views.decorators import staff_member_required
import random

def lottery_result(request, lottery_id):
    lottery = get_object_or_404(Lottery, id=lottery_id)

    # Tous les participants actifs
    participants = lottery.participations.filter(is_active=True)

    # Liste des gagnants triés par rang croissant (ordre de sélection)
    winners = list(participants.filter(is_winner=True).order_by('winner_rank'))

    # Seuls les admins peuvent voir le bouton de tirage
    show_pick_button = request.user.is_staff and participants.exclude(is_winner=True).exists()

    if request.method == "POST" and show_pick_button:
        num_to_draw = int(request.POST.get("num_winners", 1))
        available_participants = list(participants.exclude(is_winner=True))

        drawn = random.sample(available_participants, min(num_to_draw, len(available_participants)))

        # Récupérer le rang actuel maximal parmi les gagnants existants
        existing_ranks = [w.winner_rank for w in winners if w.winner_rank]
        current_rank = max(existing_ranks) if existing_ranks else 0

        for p in drawn:
            current_rank += 1
            p.is_winner = True
            p.winner_rank = current_rank
            p.save()
            winners.append(p)  # Ajouter à la liste pour affichage

        return HttpResponseRedirect(request.path)

    return render(request, 'core/lottery_result.html', {
        'lottery': lottery,
        'winners': winners,
        'participants': participants,
        'show_pick_button': show_pick_button,
    })

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import Lottery,LotteryParticipation
from .forms import LotteryParticipationForm
@login_required
def participate_in_lottery(request, lottery_id):
    lottery = get_object_or_404(Lottery, id=lottery_id, is_active=True)

    if lottery.current_participant_count() >= lottery.max_participants:
        messages.warning(request, "Le nombre maximum de participants a été atteint.")
        return redirect('lottery_list')

    if request.method == 'POST':
        form = LotteryParticipationForm(request.POST, user=request.user, lottery=lottery)
        if form.is_valid():
            form.save()
            messages.success(request, "Votre participation a été enregistrée. Elle sera activée par un administrateur.")
            return redirect('lottery_list')
    else:
        form = LotteryParticipationForm(user=request.user, lottery=lottery)

    return render(request, 'core/participate.html', {
        'form': form,
        'lottery': lottery,
    })


from django.shortcuts import render
from .models import Lottery
def lottery_list(request):
    lotteries = Lottery.objects.filter(is_active=True).order_by('-created_at')

    for lottery in lotteries:
        lottery.current_count = lottery.current_participant_count()
        # Associer le gagnant de rang 1 à cette loterie
        lottery.top_winner = (
            lottery.participations
            .filter(winner_rank=1)
            .select_related('user')
            .first()
        )
    paginator = Paginator(lotteries, 4)  # 4 tirages par page
    page = request.GET.get('page')

    try:
        lotteries = paginator.page(page)
    except PageNotAnInteger:
        lotteries = paginator.page(1)
    except EmptyPage:
        lotteries = paginator.page(paginator.num_pages)
    return render(request, 'core/lottery_list.html', {
        'lotteries': lotteries
    })

# @login_required
# def user_dashboard(request):
#     # Récupérer les points de l'utilisateur
#     try:
#         user_points = UserPoints.objects.get(user=request.user)
#     except UserPoints.DoesNotExist:
#         user_points = None

#     # Récupérer les actions effectuées par l'utilisateur
#     likes = AdInteraction.objects.filter(user=request.user, interaction_type='like').count()
#     comments = AdInteraction.objects.filter(user=request.user, interaction_type='comment').count()
#     shares = AdInteraction.objects.filter(user=request.user, interaction_type='share').count()

#     # Calcul des points pour chaque type d'action
#     likes_points = likes  # 1 point par like
#     comments_points = comments * 2  # 2 points par commentaire
#     shares_points = shares * 5  # 5 points par partage

#     # 🔥 **Ne pas recalculer `purchases_points`, on utilise directement `user_points.points`**
#     if user_points:
#         total_points = user_points.points  # ✅ Points restants après dépenses
#         spent_points = user_points.spent_points  # ✅ Ajout des points dépensés
#     else:
#         total_points = 0
#         spent_points = 0

#     return render(request, 'core/user_dashboard.html', {
#         'user_points': user_points,
#         'likes_points': likes_points,
#         'comments_points': comments_points,
#         'shares_points': shares_points,
#         'total_points': total_points,  # Points disponibles
#         'spent_points': spent_points,  # ✅ Ajout des points dépensés
#     })



# from django.http import JsonResponse
# import json

# def cart(request):
    
#     cart = None
#     cartitems = []
    
#     if request.user.is_authenticated:
#         cart, created = Cart.objects.get_or_create(user=request.user, completed=False)
#         cartitems = cart.cartitems.all()
    
#     context = {"cart":cart, "items":cartitems}
#     return render(request, "core/cart.html", context)

# # Voir la vue 'add_to_cart' existante
# @login_required
# def add_to_cart(request, product_id):
#     product = get_object_or_404(Product, id=product_id)
#     cart = get_or_create_cart(request.user)

#     # Vérifier si l'article est déjà dans le panier
#     cart_item, created = CartItem.objects.get_or_create(cart=cart, product=product)
#     if not created:
#         cart_item.quantity += 1
#         cart_item.save()

#     return redirect('core:cart_detail')

# # Ajoutez ceci dans la vue pour obtenir le nombre total d'articles
# def get_cart_item_count(user):
#     cart = get_or_create_cart(user)
#     return cart.get_item_count()
